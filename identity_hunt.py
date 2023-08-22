#!/usr/bin/python
# coding: utf-8

# <<<<<<<<<<<<<<<<<<<<<<<<<<     Copyright        >>>>>>>>>>>>>>>>>>>>>>>>>>

# Copyright (C) 2023 LincolnLandForensics
#
# This program is free software; you can redistribute it and/or modify it under
# the terms of the GNU General Public License version 2, as published by the
# Free Software Foundation
#
# This program is distributed in the hope that it will be useful, but WITHOUT
# ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
# FOR A PARTICULAR PURPOSE. See the GNU General Public License for more
# details (http://www.gnu.org/licenses/gpl.txt).

# <<<<<<<<<<<<<<<<<<<<<<<<<<     Imports        >>>>>>>>>>>>>>>>>>>>>>>>>>

try:
    from bs4 import BeautifulSoup
    import xlsxwriter
except:
    print(f'install missing modules:    pip install -r requirements_identity_hunt.txt')
    exit()

import os
import re
import sys
import json
import time
# import random
import socket
import openpyxl
import requests
import datetime
import argparse  # for menu system
from subprocess import call
from tkinter import * 
from tkinter import messagebox

from docx import Document

# <<<<<<<<<<<<<<<<<<<<<<<<<<     Pre-Sets       >>>>>>>>>>>>>>>>>>>>>>>>>>

author = 'LincolnLandForensics'
description = "OSINT: track people down by username, email, ip, phone and website"
tech = 'LincolnLandForensics'  # change this to your name if you are using Linux
version = '2.9.1'

# Regex section
# regex_host = re.compile(r'\b((?:(?!-)[a-zA-Z0-9-]{1,63}(?<!-)\.)+(?i)(?!exe|php|dll|doc' \
                        # '|docx|txt|rtf|odt|xls|xlsx|ppt|pptx|bin|pcap|ioc|pdf|mdb|asp|html|xml|jpg|gif$|png' \
                        # '|lnk|log|vbs|lco|bat|shell|quit|pdb|vbp|bdoda|bsspx|save|cpl|wav|tmp|close|ico|ini' \
                        # '|sleep|run|dat$|scr|jar|jxr|apt|w32|css|js|xpi|class|apk|rar|zip|hlp|cpp|crl' \
                        # '|cfg|cer|plg|lxdns|cgi|xn$)(?:xn--[a-zA-Z0-9]{2,22}|[a-zA-Z]{2,13}))(?:\s|$)')

regex_host = re.compile(
    r'(?i)\b((?:(?!-)[a-zA-Z0-9-]{1,63}(?<!-)\.)+'
    '(?!exe|php|dll|doc|docx|txt|rtf|odt|xls|xlsx|ppt|pptx|bin|pcap|ioc|pdf|mdb|asp|html|xml|jpg|gif$|png'
    '|lnk|log|vbs|lco|bat|shell|quit|pdb|vbp|bdoda|bsspx|save|cpl|wav|tmp|close|ico|ini'
    '|sleep|run|dat$|scr|jar|jxr|apt|w32|css|js|xpi|class|apk|rar|zip|hlp|cpp|crl'
    '|cfg|cer|plg|lxdns|cgi|xn$)'
    '(?:xn--[a-zA-Z0-9]{2,22}|[a-zA-Z]{2,13}))(?:\s|$)')


regex_md5 = re.compile(r'^([a-fA-F\d]{32})$')  # regex_md5        [a-f0-9]{32}$/gm
regex_sha1 = re.compile(r'^([a-fA-F\d]{40})$')  # regex_sha1
regex_sha256 = re.compile(r'^([a-fA-F\d]{64})$')  # regex_sha256
regex_sha512 = re.compile(r'^([a-fA-F\d]{128})$')  # regex_sha512

regex_number = re.compile(r'^(^\d)$')  # regex_number    #Beta
regex_number_fb = re.compile(r'^\d{9,15}$')  # regex_number    #to match facebook user id

regex_ipv4 = re.compile('(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}' +
                        '(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)')
regex_ipv6 = re.compile('(S*([0-9a-fA-F]{1,4}:){7,7}[0-9a-fA-F]{1,4}S*|S*(' +
                        '[0-9a-fA-F]{1,4}:){1,7}:S*|S*([0-9a-fA-F]{1,4}:)' +
                        '{1,6}:[0-9a-fA-F]{1,4}S*|S*([0-9a-fA-F]{1,4}:)' +
                        '{1,5}(:[0-9a-fA-F]{1,4}){1,2}S*|S*([0-9a-fA-F]' +
                        '{1,4}:){1,4}(:[0-9a-fA-F]{1,4}){1,3}S*|S*(' +
                        '[0-9a-fA-F]{1,4}:){1,3}(:[0-9a-fA-F]{1,4}){1,4}S*' +
                        '|S*([0-9a-fA-F]{1,4}:){1,2}(:[0-9a-fA-F]{1,4})' +
                        '{1,5}S*|S*[0-9a-fA-F]{1,4}:((:[0-9a-fA-F]{1,4})' +
                        '{1,6})S*|S*:((:[0-9a-fA-F]{1,4}){1,7}|:)S*|::(ffff' +
                        '(:0{1,4}){0,1}:){0,1}((25[0-5]|(2[0-4]|1{0,1}' +
                        '[0-9]){0,1}[0-9]).){3,3}(25[0-5]|(2[0-4]|1{0,1}[' +
                        '0-9]){0,1}[0-9])|([0-9a-fA-F]{1,4}:){1,4}:((25[' +
                        '0-5]|(2[0-4]|1{0,1}[0-9]){0,1}[0-9]).){3,3}(25[' +
                        '0-5]|(2[0-4]|1{0,1}[0-9]){0,1}[0-9]))')


regex_phone = re.compile(
    r'^(?:(?:\+?1\s*(?:[.-]\s*)?)?(?:\(\s*([2-9][0-8][0-9])\s*\)|([2-9][0-8][0-9]))\s*(?:[.-]\s*)?)?'
    r'([2-9][0-9]{2})\s*(?:[.-]\s*)?([0-9]{4})$|^(\d{10})$|^1\d{10}$')
regex_phone11 = re.compile(r'^1\d{10}$')
regex_phone2 = re.compile(r'(\d{3}) \W* (\d{3}) \W* (\d{4}) \W* (\d*)$')

# Colorize section
global color_red
global color_yellow
global color_green
global color_blue
global color_purple
global color_reset
color_red = ''
color_yellow = ''
color_green = ''
color_blue = ''
color_purple = ''
color_reset = ''

if sys.version_info > (3, 7, 9) and os.name == "nt":
    version_info = os.sys.getwindowsversion()
    major_version = version_info.major
    build_version = version_info.build

    # print(f'major version = {major_version} Build= {build_version} {version_info}')   # temp

    if major_version >= 10 and build_version >= 22000: # Windows 11 and above
        # print(f'major version = {major_version}')   # temp
        import colorama
        from colorama import Fore, Back, Style  
        print(f'{Back.BLACK}') # make sure background is black
        color_red = Fore.RED
        color_yellow = Fore.YELLOW
        color_green = Fore.GREEN
        color_blue = Fore.BLUE
        color_purple = Fore.MAGENTA
        color_reset = Style.RESET_ALL


   
# <<<<<<<<<<<<<<<<<<<<<<<<<<     Menu           >>>>>>>>>>>>>>>>>>>>>>>>>>

def main():

    # check internet status
    status = internet()
    status2 = is_running_in_virtual_machine()

    if status2 == True:
        print(f'{color_yellow}This is a virtual machine. Not checking for internet connectivity{color_reset}')
        # apparently when running from a VM (and maybe behind a proxy) it says internet isn't connected
    elif status == False:
        noInternetMsg()
        input(f'{color_red}CONNECT TO THE INTERNET FIRST. Hit Enter to exit...{color_reset}')
        exit()
    else:
        print(color_green + '\nINTERNET IS CONNECTED\n' + color_reset)

    # global section
    global filename
    filename = 'input.txt'
    global Spreadsheet
    Spreadsheet = 'Intel_.xlsx'
    # global inputDetails
    # inputDetails = 'no'

    global row
    row = 1

    global emails
    global ips
    global phones
    global users
    global dnsdomains
    global websites
    
    emails = []
    ips = []
    phones = []
    users = []
    dnsdomains = []
    websites = [] 
    
    parser = argparse.ArgumentParser(description=description)
    parser.add_argument('-I', '--input', help='', required=False)
    parser.add_argument('-O', '--output', help='', required=False)
    parser.add_argument('-E','--emailmodules', help='email modules', required=False, action='store_true')
    parser.add_argument('-b','--blurb', help='write ossint blurb', required=False, action='store_true')
    parser.add_argument('-H','--howto', help='help module', required=False, action='store_true')
    parser.add_argument('-i','--ips', help='ip modules', required=False, action='store_true')
    parser.add_argument('-p','--phonestuff', help='phone modules', required=False, action='store_true')
    parser.add_argument('-s','--samples', help='print sample inputs', required=False, action='store_true')
    parser.add_argument('-t','--test', help='testing individual modules', required=False, action='store_true')
    parser.add_argument('-U','--usersmodules', help='username modules', required=False, action='store_true')
    parser.add_argument('-w','--websitetitle', help='websites titles', required=False, action='store_true')    

    parser.add_argument('-W','--websites', help='websites modules', required=False, action='store_true')    
    
    args = parser.parse_args()
    cls()
    print_logo()

    if args.samples:  
        samples()
        return 0

    if args.blurb:
        write_blurb()

    if args.input:
        filename = args.input
    if args.output:
        Spreadsheet = args.output        

    create_ossint_xlsx()    # create the spreadsheet    
    master()

    # Check if no arguments are entered
    if len(sys.argv) == 1:
        print(f"{color_yellow}You didn't select any options so I'll run the major options{color_reset}")
        print(f'{color_yellow}try -h for a listing of all menu options{color_reset}')
        args.emailmodules = True
        args.ips = True
        args.phonestuff = True
        args.usersmodules = True
        args.websites = True
    
    if args.howto:  # this section might be redundant
        parser.print_help()
        usage()
        return 0

    if args.emailmodules:  
        ghunt()
        holehe_email()
        osintIndustries_email()
        thatsthememail()    # https://thatsthem.com/email/smooth8101@yahoo.com
        # twitteremail()    # auth required    
        wordpresssearchemail()  # requires auth
        
    if args.ips:  
        # geoiptool() # works but need need to rate limit; expired certificate breaks this
        resolverRS()
        thatsthemip()
        whoisip()   
        whatismyip()
        
    # phone modules
    if args.phonestuff:
        familytreephone()
        thatsthemphone()   # retest
        reversephonecheck()
        spydialer()
        validnumber()
        whitepagesphone()
        whocalld()
        
    if args.test:  
        print('testing')
        telegram()
        
    if args.usersmodules:  
        about()
        bitbucket() # add fullname
        blogspot_users()
        disqus()    # test
        # ebay()  # all false positives due to captcha
        etsy()
        facebook()  # works
        flickr()   # add photo, note, name, info
        freelancer()
        friendfinder()
        foursquare()    # works
        garmin()
        gravatar()
        imageshack()    # works
        instagram()
        instructables() # works
        keybase()   # works
        kik()   
        massageanywhere()
        mastadon() 
        myshopify()
        myspace_users()
        paypal()  # needs work
        patreon()
        pinterest() # works
        poshmark()    
        public()    
        snapchat()    # must manually verify
        spotify()   # works
        # telegram()# crashes the script
        threads()        
        tiktok()
        tinder() # add dob, schools
        truthSocial()
        # twitter()   # needs auth
        wordpress() # works
        wordpressprofiles()  
        youtube()   # works
        familytree()
        sherlock()
        whatsmyname()

    if args.websitetitle:  
        titles()    # alpha
        
    if args.websites:  
        redirect_detect()
        robtex()
        titles()    # alpha
        viewdnsdomain()
        whoiswebsite()    # works

    # set linux ownership    
    if sys.platform == 'win32' or sys.platform == 'win64':
        pass
    else:
        call(["chown %s.%s *.xlsx" % (tech.lower(), tech.lower())], shell=True)

    workbook.close()
    if not args.blurb:
        input(f"See '{Spreadsheet}' for output. Hit Enter to exit...")

    return 0
    
    # sys.exit()  # this code is unreachable


# <<<<<<<<<<<<<<<<<<<<<<<<<<  Sub-Routines   >>>>>>>>>>>>>>>>>>>>>>>>>>

def cls():
    linux = 'clear'
    windows = 'cls'
    os.system([linux, windows][os.name == 'nt'])

def master():
    global row  # The magic to pass row globally
    style = workbook.add_format()
    color = 'white'
    if not os.path.exists(filename):
        input(f"{color_red}{filename} doesnt exist.{color_reset}")
        sys.exit()
    elif os.path.getsize(filename) == 0:
        input(f'{color_red}{filename} is empty. Fill it with username, email, ip, phone and/or websites.{color_reset}')
        sys.exit()
    elif os.path.isfile(filename):
        print(f'{color_green}Reading {filename}{color_reset}')        
        inputfile = open(filename)
    else:
        input(f'{color_red}See {filename} does not exist. Hit Enter to exit...{color_reset}')
        sys.exit()
        
    for eachline in inputfile:
        
        (query, ranking, fullname, url, email , user) = ('', '', '', '', '', '')
        (phone, ip, entity, fulladdress, city, state) = ('', '', '', '', '', '')
        # (ipv6) = ('')
        (zip, country, note, aka, dob, gender) = ('', '', '', '', '', '')
        (info, misc, lastname, firstname, middlename, friend) = ('', '', '', '', '', '')
        (otherurls, otherphones, otheremails, case, sosfilenumber, president) = ('', '', '', '', '', '')
        (sosagent, managers, dnsdomain, dstip, srcip  ) = ('', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        # (ips, emails, phones) = ('', '', '') # test
        
        (color) = 'white'
        style.set_bg_color('white')  # test

        eachline = eachline + "\t" * 40
        eachline = eachline.split('\t')  # splits by tabs

        query = (eachline[0].strip())
        ranking = (eachline[1].strip())
        if ranking == '':
            ranking = '1 - main'
        fullname = (eachline[2].strip())
        url = (eachline[3].strip())
        email = (eachline[4].strip())
        user = (eachline[5].strip())
        phone = (eachline[6].strip())
        ip = (eachline[7].strip())
        entity = (eachline[8].strip())
        fulladdress = (eachline[9].strip())
        city = (eachline[10].strip())
        state = (eachline[11].strip())
        zip = (eachline[12].strip())
        country = (eachline[13].strip())
        note = (eachline[14].strip())
        aka = (eachline[15].strip())
        dob = (eachline[16].strip())
        gender = (eachline[17].strip())
        info = (eachline[18].strip())
        misc = (eachline[19].strip())
        lastname = (eachline[20].strip())
        firstname = (eachline[21].strip())
        middlename = (eachline[22].strip())
        friend = (eachline[23].strip())
        otherurls = (eachline[24].strip())
        otherphones = (eachline[25].strip())
        otheremails = (eachline[26].strip())
        case = (eachline[27].strip())
        sosfilenumber = (eachline[28].strip())
        president = (eachline[29].strip())
        sosagent = (eachline[30].strip())
        managers = (eachline[31].strip())
        dnsdomain = (eachline[32].strip())
        dstip = (eachline[33].strip())
        srcip = (eachline[34].strip())
        content = (eachline[35].strip())
        referer = (eachline[36].strip())
        osurl = (eachline[37].strip())
        titleurl = (eachline[38].strip())
        pagestatus = (eachline[39].strip())        

        # Regex data type
        if bool(re.search(r"^[\w\.\+\-]+\@[\w]+\.[a-z]{2,3}$", query)):  # regex email    
            email = query
            user = email.split('@')[0]
            temp1 = [email]
            if query.lower() not in emails:            # don't add duplicates
                emails.append(email)
        elif re.search(regex_host, query):  # regex_host (Finds url and dnsdomain) # False positives for emails    # todo
            url = query
            if url.lower().startswith('http'):
                if url.lower() not in websites:            # don't add duplicates
                    websites.append(url)            
            else:
                logsource = 'IOC-dnsdomain'
                dnsdomain = query
            url = url.rstrip('/')
            if url.lower() not in websites:            # don't add duplicates
                websites.append(url)            
            dnsdomain = url.lower()
            dnsdomain = dnsdomain.replace("https://", "")
            dnsdomain = dnsdomain.replace("http://", "")
            dnsdomain = dnsdomain.split('/')[0]
            notes2 = dnsdomain.split('.')[-1]
            if dnsdomain.lower() not in dnsdomains:            # don't add duplicates
                dnsdomains.append(dnsdomain)
            
        elif re.search(regex_ipv4, query):  # regex_ipv4
            (ip) = (query)
            if query.lower() not in ips:            # don't add duplicates
                ips.append(ip)

        elif re.search(regex_ipv6, query):  # regex_ipv6
            (ip) = (query)
            if query.lower() not in ips:            # don't add duplicates
                ips.append(ip)

        elif re.search(regex_phone, query) or re.search(regex_phone11, query) or re.search(regex_phone2, query):  # regex_phone
            (phone) = (query)
            if query.lower() not in phones:            # don't add duplicates
                phones.append(phone)

            
        elif query.lower().startswith('http'):
            url = query
            if url.lower() not in websites:            # don't add duplicates
                websites.append(url)            
        elif query.strip() == '':
            print(f'{color_red}blank input found{color_reset}')
        else:
            user = query
            if query.lower() not in users:            # don't add duplicates
                users.append(user)
            # (srcip, dstip) = (note, note)
            # logsource = 'IOC-IP'
        # elif re.search(regex_md5, note):  # regex_md5
            # logsource = 'IOC-Hash-Md5'
        # elif re.search(regex_sha256, note):  # regex_sha256
            # logsource = 'IOC-Hash-Sha256'
        # elif re.search(regex_sha1, note):  # regex_sha1
            # logsource = 'IOC-Hash-Sha1'
        # elif re.search(regex_sha512, note):  # regex_sha512
            # logsource = 'IOC-Hash-Sha512'

        # Read url,dnsdomain or dstip
        # if url != '':
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
        # elif dnsdomain != '':
            # (content, referer, osurl, titleurl, pagestatus) = request(dnsdomain)
        # elif dstip != '':
            # (content, referer, osurl, titleurl, pagestatus) = request(dstip)
        # copy whole page
        # payload = content  # temp

        # color
        # if 'yahoo' in url:  
            # format_function(bg_color='orange')
        # elif 'google' in url:  #
            # format_function(bg_color='green')
        # elif 'Fail' in pagestatus:  #
            # format_function(bg_color='red')
        # else:
            # format_function(bg_color='white')

        # Write OSSINT excel
        write_ossint(query, ranking, fullname, url, email, user, phone, ip, entity, 
            fulladdress, city, state, zip, country, note, aka, dob, gender, info, 
            misc, lastname, firstname, middlename, friend, otherurls, otherphones, 
            otheremails, case, sosfilenumber, president, sosagent, managers, 
            dnsdomain, dstip, srcip, content, referer, osurl,
            titleurl, pagestatus)
    return users,ips,emails,phones


def about(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< about.me {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (fulladdress, city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '', '')
        (lastname, firstname) = ('', '')
        user = user.rstrip()
        url = ('https://about.me/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        for eachline in content.split("\n"):
            if " | about.me" in eachline:        
                fullname = eachline.strip().replace(" | about.me","") # .split("-")(0)

                if ' - ' in fullname:
                    fulladdress = fullname.split(' - ')[1]
                    fullname = fullname.split(' - ')[0]  
                    if ' ' in fullname:
                        fullname2 = fullname.split(' ')

                        firstname = fullname2[0]
                        lastname = fullname2[1]   
                    

        if '404' not in pagestatus:
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '3 - about.me', fullname, url, '', user, '', '', '', fulladdress
                , city, '', '', country, '', '', '', '', '', '', lastname, firstname, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        

def fastpeoplesearch():# testPhone= 385-347-1531
    print(f'{color_yellow}\n\t<<<<< fastpeoplesearch {color_blue}phone numbers{color_yellow} >>>>>{color_reset}')
    
    for phone in phones:
        (country, city, zip, case, note, content) = ('', '', '', '', '', '')
        (fullname, content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '', 'research')
        phone = phone.replace('(','').replace(')','-').replace(' ','')

        url = ('https://www.fastpeoplesearch.com/%s' %(phone))
        # (content, referer, osurl, titleurl, pagestatus) = request(url)    # protected by cloudflare

        for eachline in content.split("\n"):
            if "We could not find any results based on your search criteria" in eachline and case == '':
                print(f'{color_red}Not found{color_reset}')  # temp
                url = ('')
      
                
        if url != '':
            print(f'{url}') 
            write_ossint(phone, '9 - fastpeoplesearch', fullname, url, '', '', phone, '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)

def blogspot_users(): # testuser = kevinrose
    print(f'\n\t<<<<< blogspot users >>>>>')
    
    for user in users:
        url = f"https://{user}.blogspot.com"

        (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        (fullname) = ('')

        if 'Success' in pagestatus:
            titleurl = titleurl_og(content)
            fullname = titleurl

            write_ossint(user, '4 - blogspot', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)


def bitbucket(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< bitbucket {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = ('https://bitbucket.org/%s/' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        # for eachline in content.split("\n"):
            # if "display_name" in eachline:
                # fullname = eachline
        
        if '404' not in pagestatus:
            # grab display_name = fullname
            titleurl = titleurl.replace('Bitbucket','')
            print(f'{color_green}{url}{color_yellow}	{titleurl}{color_reset}') 
            write_ossint(user, '4 - bitbucket', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '')   

                
def create_ossint_xlsx():
    global workbook
    workbook = xlsxwriter.Workbook(Spreadsheet)
    global sheet1
    sheet1 = workbook.add_worksheet('intel')

    headerformat = workbook.add_format({'bold': True, 'border': 1})
    sheet1.freeze_panes(1, 1)  # Freeze cells
    sheet1.set_selection('B2')


    # Excel column width
    sheet1.set_column(0, 0, 23) # query
    sheet1.set_column(1, 1, 14) # ranking
    sheet1.set_column(2, 2, 16) # fullname
    sheet1.set_column(3, 3, 25) # url
    sheet1.set_column(4, 4, 23) # email
    sheet1.set_column(5, 5, 14) # user
    sheet1.set_column(6, 6, 16) # phone
    sheet1.set_column(7, 7, 15) # ip
    sheet1.set_column(8, 8, 16) # entity
    sheet1.set_column(9, 9, 25) # fulladdress
    sheet1.set_column(10, 10, 15) # city
    sheet1.set_column(11, 11, 5) # state
    sheet1.set_column(12, 12, 7) # zip
    sheet1.set_column(13, 13, 8) # country
    sheet1.set_column(14, 14, 16) # note
    sheet1.set_column(15, 15, 15) # aka
    sheet1.set_column(16, 16, 12) # dob
    sheet1.set_column(17, 17, 7) # gender
    sheet1.set_column(18, 18, 15) # info
    sheet1.set_column(19, 19, 8) # misc
    sheet1.set_column(20, 20, 10) # lastname
    sheet1.set_column(21, 21, 10) # firstname
    sheet1.set_column(22, 22, 8) # middlename
    sheet1.set_column(23, 23, 17) # friend
    sheet1.set_column(24, 24, 15) # otherurls
    sheet1.set_column(25, 25, 15) # otherphones
    sheet1.set_column(26, 26, 15) # otheremails
    sheet1.set_column(27, 27, 12) # case
    sheet1.set_column(28, 28, 9) # sosfilenumber
    sheet1.set_column(29, 29, 12) # president
    sheet1.set_column(30, 30, 12) # sosagent
    sheet1.set_column(31, 31, 12) # managers
    sheet1.set_column(32, 32, 16) # dnsdomain
    sheet1.set_column(33, 33, 12) # dstip
    sheet1.set_column(34, 34, 12) # srcip
    sheet1.set_column(35, 35, 9) # content
    sheet1.set_column(36, 36, 9) # referer
    sheet1.set_column(37, 37, 6) # osurl
    sheet1.set_column(38, 38, 10) # titleurl
    sheet1.set_column(39, 39, 12) # pagestatus



    # Write column headers

    sheet1.write(0, 0, 'query', headerformat)
    sheet1.write(0, 1, 'ranking', headerformat)
    sheet1.write(0, 2, 'fullname', headerformat)
    sheet1.write(0, 3, 'url', headerformat)
    sheet1.write(0, 4, 'email', headerformat)
    sheet1.write(0, 5, 'user', headerformat)
    sheet1.write(0, 6, 'phone', headerformat)
    sheet1.write(0, 7, 'ip', headerformat)
    sheet1.write(0, 8, 'business/entity', headerformat)
    sheet1.write(0, 9, 'fulladdress', headerformat)
    sheet1.write(0, 10, 'city', headerformat)
    sheet1.write(0, 11, 'state', headerformat)
    sheet1.write(0, 12, 'zip', headerformat)
    sheet1.write(0, 13, 'country', headerformat)
    sheet1.write(0, 14, 'note', headerformat)
    sheet1.write(0, 15, 'aka', headerformat)
    sheet1.write(0, 16, 'dob', headerformat)
    sheet1.write(0, 17, 'gender', headerformat)
    sheet1.write(0, 18, 'info', headerformat)
    sheet1.write(0, 19, 'misc', headerformat)
    sheet1.write(0, 20, 'lastname', headerformat)
    sheet1.write(0, 21, 'firstname', headerformat)
    sheet1.write(0, 22, 'middlename', headerformat)
    sheet1.write(0, 23, 'friend', headerformat)
    sheet1.write(0, 24, 'otherurls', headerformat)
    sheet1.write(0, 25, 'otherphones', headerformat)
    sheet1.write(0, 26, 'otheremails', headerformat)
    sheet1.write(0, 27, 'case', headerformat)
    sheet1.write(0, 28, 'sosfilenumber', headerformat)
    sheet1.write(0, 29, 'president', headerformat)
    sheet1.write(0, 30, 'sosagent', headerformat)
    sheet1.write(0, 31, 'managers', headerformat)
    sheet1.write(0, 32, 'dnsdomain', headerformat)
    sheet1.write(0, 33, 'dstip', headerformat)
    sheet1.write(0, 34, 'srcip', headerformat)
    sheet1.write(0, 35, 'content', headerformat)
    sheet1.write(0, 36, 'referer', headerformat)
    sheet1.write(0, 37, 'osurl', headerformat)
    sheet1.write(0, 38, 'titleurl', headerformat)
    sheet1.write(0, 39, 'pagestatus', headerformat)

    # hidden columns
    sheet1.set_column(33, 33, None, None, {'hidden': 1}) # dstip
    sheet1.set_column(34, 34, None, None, {'hidden': 1}) # srcip

    # sheet 2

    global sheet2    
    sheet2 = workbook.add_worksheet('color code')

    sheet2.freeze_panes(1, 1)  # Freeze cells
    sheet2.set_selection('B2')

    sheet2.set_column(0, 0, 8) # color
    sheet2.set_column(1, 1, 23) # Color code description


    # Define formats for different cell colors
    red_format = workbook.add_format({'bg_color': 'red'})   # red
    green_format = workbook.add_format({'bg_color': '#92D050'}) # green
    orange_format = workbook.add_format({'bg_color': '#FFc000'}) # orange

    sheet2.write(0, 0, 'color', headerformat)
    sheet2.write(0, 1, 'color description', headerformat)
    sheet2.write(1, 0, 'green', green_format)
    sheet2.write(2, 0, 'orange', orange_format)
    sheet2.write(3, 0, 'red', red_format)


    sheet2.write(1, 1, 'Verified or High Confidence', green_format)
    sheet2.write(2, 1, 'Research/Verify', orange_format)
    sheet2.write(3, 1, 'False Positive or Dead Link', red_format)





def disqus(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< discus {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = ('http://disqus.com/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        if '404' not in pagestatus:
            # fullname = titleurl
            print(f'{color_green}{url}{color_reset}') 
            write_ossint(user, '5 - discus', '', url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)        


def ebay(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< ebay {color_blue}users{color_yellow} >>>>>{color_reset}')
    print(f'{color_yellow}\n\tthis can take a while >>>>>{color_reset}')

    for user in users:    
        (city, country, fullname, titleurl, pagestatus, note) = ('', '', '', '', '', '')
        user = user.rstrip()
        url = ('https://www.ebay.com/usr/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        for eachline in content.split("\n"):
            if 'been an eBay member since' in eachline:
                note = ('%s %s' %(eachline.strip(), note))
        
        if 'Positive feedback' in content:
            fullname = titleurl
            print(f'{color_green}{url}{color_reset}{titleurl}') 
            write_ossint(user, '7 - ebay', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)        
        time.sleep(5) #will sleep for 5 seconds
          
def etsy(): # testuser = kevinrose https://www.etsy.com/people/kevinrose
    print(f'{color_yellow}\n\t<<<<< etsy {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = ('https://www.etsy.com/people/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        if '404' not in pagestatus:
            # grab display_name = fullname
            titleurl = titleurl.replace("'s favorite items - Etsy",'')
            print(f'{color_green}{url}{color_yellow}	{titleurl}{color_reset}') 
            write_ossint(user, '4 - etsy', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)   

def facebook(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< facebook {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (Success,FullName,LastName,FirstName,ID,Gender) = ('','','','','','')
        (Photo,Country,Website,Email,Language,Username) = ('','','','','','')
        (city, country) = ('', '')
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        user = user.rstrip()
        url = ('http://facebook.com/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        # try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
        # except:
            # pass

        if 'Success' in pagestatus and 'vi-vn.facebook.com' in content:
            fullname = titleurl
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '3 - Facebook', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)

    for phone in phones:
        (country, city, state, zip, case, note) = ('', '', '', '', '', '')
        (fullname, content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '', '')
        phone = phone.replace('(','').replace(')','').replace(' ','')
        
        if re.search(regex_number_fb, phone):
            url = ('https://www.facebook.com/%s' %(phone))
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            if 'Success' in pagestatus and 'vi-vn.facebook.com' in content:
                user = phone
                phone = ''
                fullname = titleurl

            if 1==1:  
                print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}')
                write_ossint(phone, '8 - facebook', fullname, url, '', user, '', '', '', ''
                    , city, state, '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, pagestatus)


def familytree(): 

    print(f'\n\t{color_yellow}<<<<< Manually check familytreenow.com >>>>>{color_reset}')
    url = ('https://www.familytreenow.com/search/')
    write_ossint('', '9 - manual', '', url, '', '', '', '', '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '')

def familytreephone():# testPhone= 708-372-8101 DROP THE LEADING 1
    print(f'{color_yellow}\n\t<<<<< familytree {color_blue}phone numbers{color_yellow} >>>>>{color_reset}')
    for phone in phones:
        (country, city, state, zip, case, note) = ('', '', '', '', '', '')
        (fullname, content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '', '')
        phone = phone.replace('(','').replace(')','').replace(' ','')
        
        if phone.startswith('1'):
            phone = phone.replace('1','')
        url = ('https://www.familytreenow.com/search/genealogy/results?phoneno=%s' %(phone.replace("-", "")))
      
        if 1==1:        
            print(f'{color_yellow}{url}{color_reset}')
            write_ossint(phone, '8 - familytree', fullname, url, '', '', phone, '', '', ''
                , city, state, '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, pagestatus)


def flickr(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< flickr {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.rstrip()
        url = ('https://www.flickr.com/people/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("  <"):
            if "og:title" in eachline:
                fullname = eachline.strip().split("\"")[1]

        if '404' not in pagestatus and 'ail' not in pagestatus:
            if fullname.lower() == user.lower():
                fullname = ''

            titleurl = titleurl.replace('About ','').replace(" | Flickr",'')
            print(f'{color_green}{url}{color_yellow}	{titleurl}{color_reset}') 
            write_ossint(user, '4 - flickr', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)        

           
def format_function(bg_color='white'):
    global format
    format = workbook.add_format({
        'bg_color': bg_color
    })


def foursquare():    # testuser=    john
    print(f'{color_yellow}\n\t<<<<< foursquare {color_blue}users{color_yellow} >>>>>{color_reset}')    

    for user in users:    
        url = ('https://foursquare.com/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)

            content = content.strip()
            titleurl = titleurl.strip()
       
        except:
            pass

        if ' on Foursquare' in titleurl:
            fullname = titleurl.rstrip(' on Foursquare')
        
            write_ossint(user, '7 - foursquare', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_reset}')    
            
            
def freelancer(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< freelancer {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = ('https://www.freelancer.com/u/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        titleurl = titleurl.replace(' Profile | Freelancer','')
        if '404' not in pagestatus:
           
            if ' ' in titleurl:
                fullname = titleurl
            
            if fullname.lower() == user.lower():
                fullname = ''
            
            
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '5 - freelancer', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        

def friendfinder():    # testuser=  kevinrose
    print(f'{color_yellow}\n\t<<<<< friendfinder {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        url = ('https://www.friendfinder-x.com/profile/%s' %(user))

        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
      
        except:
            pass

        if 'Register to Find' not in titleurl:
            write_ossint(user, '7 - friendfinder.com', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_reset}')    

def garmin(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< garmin {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (Success, fullname, lastname, firstname, case, gender) = ('','','','','','')
        (photo, country, website, email, language, username) = ('','','','','','')
        (city) = ('')
        user = user.rstrip()

        url = ('https://connect.garmin.com/modern/profile/%s' %(user))

        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        # try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
        # except:
            # pass
        # if 1==1:
        # print(content)
        # if user in content:
        if 'twitter:card' not in content:
        
            fullname = titleurl
            fullname = fullname.split(" (")[0]
            fullname = fullname.replace("Garmin Connect","").strip()
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '9 - garmin', '', url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)

 
def geoiptool():    # testuser= 77.15.67.232
    print(f"{color_yellow}\n\t<<<<< geodatatool {color_blue}IP's{color_yellow} >>>>>{color_reset}")

    for ip in ips:
        (country, city, state, zip) = ('', '', '', '')
     
        # url = ('http://geoiptool.com/en/?IP=%s' %(ip))
        url = ('https://www.geodatatool.com/en/?IP=%s' %(ip))

        # (content, referer, osurl, titleurl, pagestatus) = request(url)    # certificate errors
        (content, referer, osurl, titleurl, pagestatus) = request_url(url)  # ignore certificates
        print(type(content))    # temp
        content2 = content.strip().split("\n")
        
        # for eachline in content:        
        for eachline in content2:
        
        # for eachline in content.split("\n"):
            print(f'Eachline = {eachline}') # temp

            if "Country: " in eachline:
                country = eachline.strip().split(": ")[1]
                country = country.split("<")[0]
            elif "City: " in eachline:
                city = eachline.strip().split(": ")[1]
                city = city.split("<")[0]            
            elif "Region: " in eachline:
                state = eachline.strip().split(": ")[1]
                state = state.split("<")[0]            
            elif "Postal Code: " in eachline:
                zip = eachline.strip().split(": ")[1]
                zip = zip.split("<")[0] 
        print(f'{color_yellow}{ip}  {country}   {city}  {state} {zip}{color_reset}')
        time.sleep(7) #will sleep for 30 seconds
        write_ossint(ip, '6 - geodatatool', '', url, '', '', '', ip, '', ''
        , city, state, zip, country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)

def ghunt():# testEmail= kevinrose@gmail.com   
    for email in emails:
        (note, url) = ('', '')
        note = ('cd C:\Forensics\scripts\python\git-repo\GHunt && ghunt email %s' %(email)) 
        if email.endswith('gmail.com'):
            ranking = ('8 - ghunt')
        else:
            ranking = ('9 - ghunt')  
        write_ossint(email, ranking, '', url, email, '', '', '', '', ''
        , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', 'research')


def gravatar(): # testuser = kevinrose      https://en.gravatar.com/kevinrose
    print(f'{color_yellow}\n\t<<<<< gravatar {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        (info, lastname, firstname, note, otherurls) = ('', '','', '', '')
        user = user.rstrip()
        url = (' https://gravatar.com/%s.json' %(user))        
 
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        if 'ail' not in pagestatus:
            parsed_data = json.loads(content)

            info = parsed_data['entry'][0]['photos'][0]['value']


            if 'familyName' in content:
                firstname = parsed_data['entry'][0]['name']['givenName']
                lastname = parsed_data['entry'][0]['name']['familyName']
                fullname = parsed_data['entry'][0]['name']['formatted']
            if 'aboutMe' in content:
                note = parsed_data['entry'][0]['aboutMe']
                note = note.replace('&amp', '&')
                
            if 'urls\":[{\"value' in content and 'aboutMe' in content:

                # Extract the urls list from the JSON object
                urls = parsed_data['entry'][0]['urls']

                for url in urls:
                    otherurls = ('%s %s' %(url['value'], otherurls))

            url = ('http://en.gravatar.com/%s' %(user))
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            
            if fullname != '' or otherurls != '' or note != '': 
            
                write_ossint(user, '3 - gravatar', fullname, url, '', user, '', '', '', ''
                    , city, '', '', country, note, '', '', '', info, '', lastname, firstname, '', '', otherurls, '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        
            else:
                write_ossint(user, '7 - gravatar', fullname, url, '', user, '', '', '', ''
                    , city, '', '', country, note, '', '', '', info, '', lastname, firstname, '', '', otherurls, '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        
        

def holehe_email():# testEmail= kevinrose@gmail.com
    print(f'{color_yellow}\n\t<<<<< holehe {color_blue}emails{color_yellow} >>>>>{color_reset}')
    for email in emails:
        (country, city, zip, case) = ('', '', '', '')
        
        url = ('cd C:\Forensics\scripts\python\git-repo\holehe && holehe -NP --no-color --no-clear --only-used %s' %(email))
        write_ossint(email, '9 - manual', '', url, email, '', '', '', '', ''
            , '', '', '', '', 'https://github.com/megadose/holehe', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', 'research')

def imageshack(): # testuser = ToddGilbert
    print(f'{color_yellow}\n\t<<<<< imageshack {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        # (Success,FullName,LastName,FirstName,ID,Gender) = ('','','','','','')
        # (Photo,Country,Website,Email,Language,Username) = ('','','','','','')
        (city, country, fullname) = ('', '', '')
        user = user.rstrip()
        url = ('https://imageshack.com/user/%s' %(user))

        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except:
            pass

        if 's Images' in titleurl:
            # fullname = titleurl
            print(f'{color_green}{url}{color_reset}') 
            write_ossint(user, '4 - imageshack', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)        

def instagram():    # testuser=    kevinrose     # add info
    print(f'{color_yellow}\n\t<<<<< instagram {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        url = ('https://instagram.com/%s/' %(user))
        # https://i.instagram.com/api/v1/users/<profile_id>/info/
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            
            
            
            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    fullname = eachline.split('\"')[1].split(' (')[0]
                elif 'ProfilePage\",\"description' in eachline:
                    info = eachline
                    # Load the JSON data
                    data = json.loads(eachline)

                    # Extract the description value and print it
                    note = data['description']
       
        except:
            pass
            
        # time.sleep(1) # will sleep for 1 seconds
        if 'alternate' in content:
            
        
            write_ossint(user, '3 - instagram', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_reset}')    

 
def instagramtwo(): #alpha
    # from lib.colors import red,white,green,reset
    self = 'kevinrose'

    response = self.session.get(self.url)
    if response.status_code != 200:
        exit(f'{color_red}[-] instagram: user not found{color_reset}')
    response = response.json()
    user = response['graphql']['user']
           
    data = {'Profile photo': user['profile_pic_url_hd'],
                 'Username': user['username'],
                 'User ID': user['id'],
                 'External URL': user['external_url'],
                 'Bio': user['biography'],
                 'Followers': user['edge_followed_by']['count'],
                 'Following': user['edge_follow']['count'],
                 'Pronouns': user['pronouns'],
                 'Images': user['edge_owner_to_timeline_media']['count'],
                 'Videos': user['edge_felix_video_timeline']['count'],
                 'Reels': user['highlight_reel_count'],
                 'Is private?': user['is_private'],
                 'Is verified?': user['is_verified'],
                 'Is business account?': user['is_business_account'],
                 'Is professional account?': user['is_professional_account'],
                 'Is recently joined?': user['is_joined_recently'],
                 'Business category': user['business_category_name'],
                 'Category': user['category_enum'],
                 'Has guides?': user['has_guides']
    }
    print
    
    print(f"\n{user['full_name']} | Instagram{reset}")
    for key, value in data.items():
       print(f"├─ {key}: {value}")


def instructables(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< instructables {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = ('https://www.instructables.com/member/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        if '404' not in pagestatus:
            
            titleUrl = titleurl.replace("'s Profile - Instructables","")
            fullname = titleurl
            print(f'{color_green}{url}{color_yellow}	{titleurl}{color_reset}') 
            write_ossint(user, '7 - instructables', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        

def internet(host="8.8.8.8", port=53, timeout=3):
    """
    Host: 8.8.8.8 (google-public-dns-a.google.com)
    OpenPort: 53/tcp
    Service: domain (DNS/TCP)
    """
    try:
        socket.setdefaulttimeout(timeout)
        socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect((host, port))
        return True
    except socket.error as ex:
        print(ex)
        return False
 
            
def ip_address(dnsdomain):
    (ip) = ('')
    """
    Ping the URL and return the IP address
    """
    try:
        ip = socket.gethostbyname(dnsdomain)
    except socket.gaierror:
        ip = ''
    return ip

def is_running_in_virtual_machine():
    # Check for common virtualization artifacts
    virtualization_artifacts = [
        "/dev/virtio-ports",
        "/dev/vboxguest",
        "/dev/vmware",
        "/dev/qemu",
        "/sys/class/dmi/id/product_name",
        "/proc/scsi/scsi",
    ]

    for artifact in virtualization_artifacts:
        if os.path.exists(artifact):
            return True
            print('This is running in a virtual machine')
    return False
    
def noInternetMsg():
    '''
    prints a pop-up that says "Connect to the Internet first"
    '''
    window = Tk()
    window.geometry("1x1")
      
    w = Label(window, text ='Translate-Inator', font = "100") 
    w.pack()
    messagebox.showwarning("Warning", "Connect to the Internet first") 

def keybase():    # testuser=    kevin
    print(f'{color_yellow}\n\t<<<<< keybase.io {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        url = ('https://keybase.io/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        try:

            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    fullname = eachline.split('\"')[1].split(' (')[0]
                elif 'ProfilePage\",\"description' in eachline:
                    info = eachline
                    # Load the JSON data
                    data = json.loads(eachline)

                    # Extract the description value and print it
                    note = data['description']
       
        except:
            pass
            
        # time.sleep(1) # will sleep for 1 seconds
        # if 'what you are looking for...it does not exist' not in content:
        # if 'Your conversation will be end-to-end encrypted' in content:
        if 'Following' in content:
            write_ossint(user, '3 - keybase.com', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_reset}')    
    
def kik(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< kik {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (fullname, titleurl, pagestatus, content) = ('', '', '', '')
        (note, firstname, lastname, photo, misc, lastseen) = ('', '', '', '', '', '')
        (otherurl) = ('')
        user = user.rstrip()
        url = ('https://ws2.kik.com/user/%s' %(user))
        
        otherurl = ('http://kik.me/%s' %(user))
               
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        for eachline in content.split(","):
            if "firstName" in eachline:
                firstname = eachline.strip().split(":")[1]
                firstname = firstname.strip('\"')
            elif "lastName" in eachline:
                lastname = eachline.strip().split(":")[1]
                lastname = lastname.strip('\"')
            elif "displayPicLastModified" in eachline:
                note = eachline.strip().split(":")[1]
             
                
            elif "displayPic\"" in eachline:
                photo = eachline.strip().split(":\"")[1].split("\"")[0].replace("\\","")
                
            fullname = ('%s %s' %(firstname,lastname))
            fullname = fullname.replace("\"}","")
        if '404' not in pagestatus:
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}')
            
            write_ossint(user, '4 - kik', fullname, otherurl, '', user, '', '', '', ''
                , '', '', '', '', photo, '', '', '', '', misc, lastname, firstname, '', '', url, '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        

def massageanywhere():    # testuser=   Misty0427
    print(f'{color_yellow}\n\t<<<<< massageanywhere {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        url = ('https://www.massageanywhere.com/profile/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note, firstname, fulladdress) = ('', '', '', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
       
        except:
            pass

        if 'Profile for' in titleurl:
            if 'MassageAnywhere.com Profile for ' in titleurl:  
                titleurl = titleurl.replace('MassageAnywhere.com Profile for ','')
                if ' of ' in titleurl:
                    # titleurl = titleurl.split(' of ')
                    fullname = titleurl.split(' of ')[0]
                    fulladdress = titleurl.split(' of ')[1]
                    
            write_ossint(user, '7 - massageanywhere.com', fullname, url, '', user, '', '', '', fulladdress
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_reset}')    


def mastadon(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< mastadon {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (Success, fullname, lastname, firstname, case, gender) = ('','','','','','')
        (photo, country, website, email, language, username) = ('','','','','','')
        (city, note, info, email, content, pagestatus) = ('', '', '', '', '', '')
        (lastname, firstname, data) = ('', '', '')
        
        user = user.rstrip()
        url = ('https://mastodon.social/@%s' %(user))
        note = ('https://mastodon.social/api/v2/search?q=%s' %(user))

        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "og:title" in eachline:
                fullname = eachline.strip().split("\"")[1].split(' (')[0]

        if 'accounts\":[{' in content:        
            data = json.loads(content)              # Convert JSON data to Python dictionary
            fullname = data["accounts"][0]["display_name"]
            info = data['accounts'][0]['avatar']
        
        if " " in fullname:
            firstname = fullname.split(" ")[0]
            lastname = fullname.split(" ")[1]
            # print(data)
        if "uccess" in pagestatus and 'This resource could not be found' not in content:
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '3 - mastadon', fullname, url, email, user, '', '', '', ''
               , city, '', '', country, note, '', '', '', info, '', lastname, firstname, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)

def myshopify():    # testuser=    rothys
    print(f'{color_yellow}\n\t<<<<< myshopify {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        url = ('https://%s.myshopify.com/' %(user))
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note, otherurls) = ('', '', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    info = eachline.split('\"')[1].split(' (')[0]
                elif 'ProfilePage\",\"description' in eachline:
                    info = eachline
                    # # Load the JSON data
                    data = json.loads(eachline)

                    # # Extract the description value and print it
                    note = data['description']

        except:
            pass
            
        time.sleep(1) # will sleep for 1 seconds
        if 'Success' in pagestatus:

            response = requests.get(url)

            if response.history:
                otherurls = response.url
                note = ('redirects to %s' %(otherurls))

            write_ossint(user, '5 - myshopify.com', '', url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', info, '', '', '', '', '', otherurls, '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}	{color_yellow}{note}{color_reset}')    

def myspace_users():
    print(f'{color_yellow}\n\t<<<<< myspace {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:
        url = f"https://myspace.com/{user}"

        (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        (fullname) = ('')

        if 'Success' in pagestatus and ('Your search did not return any results') not in content:
            fullname = titleurl

            print(f'{color_green}{url}{color_yellow}	   {fullname}{color_reset}')
            write_ossint(user, '4 - myspace', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)

            
def osintIndustries_email():

    print(f'{color_yellow}\n\t<<<<< osint.Industries entry >>>>>{color_reset}')
    url = ('https://osint.industries/email#')
    write_ossint('', '9 - manual', '', url, '', '', '', '', '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '')
    
    
def patreon(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< patreon {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        (note) = ('')
        user = user.rstrip()
        url = ('https://www.patreon.com/%s/creators' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        if '404' not in pagestatus:
            titleurl = titleurl.replace(' Patreon','')
            if ' | ' in titleurl:   
                fullname = titleurl.split(' | ')[0]
                if user == fullname:
                    fullname = ''
                titleUrl = titleurl.replace("Patreon","").strip()
                note = titleurl.split(' | ')[1]
            print(f'{color_green}{url}{color_yellow}	{titleurl}{color_reset}') 
            
            
            write_ossint(user, '5 - patreon', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        


def paypal(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< paypal {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        (fulladdress, lastname, firstname) = ('', '', '')
        (email, phone, otherurl) = ('', '', '')
        user = user.rstrip()
        url = ('https://www.paypal.com/paypalme/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        (note) = ('')
        if '404' not in pagestatus:
            for eachline in content.split("\n"):
                if eachline == "": pass                                             # skip blank lines
                else:
                    # define the regular expression pattern
                    pattern = r'{"userInfo":{(.*?)}}'

                    # match the pattern to the input string
                    match = re.search(pattern, content)

                    # extract the data variable from the match object
                    if match:
                        # data = match
                        data = match.group(1)
                        note = data
                        note = note.replace("null",'\"\"')
        # else:
            # print(f'{color_red}{user}{color_reset}') 
        if ':' in note:
            titleUrl = titleurl.replace('PayPal','').strip() # task
            fullname = titleurl       

            # Extract variables using regex
            try:
                firstname = re.search(r'"givenName":"(.*?)"', data).group(1)
            except:pass
            try:
                lastname = re.search(r'"familyName":"(.*?)"', data).group(1)
            except:pass
            try:
                fullname = re.search(r'"displayName":"(.*?)"', data).group(1)
            except:pass                
            try:    
                email = re.search(r'"displayEmail":(null|".*?")', data).group(1)
            except:pass                
            try:    
                phone = re.search(r'"displayMobilePhone":(null|".*?")', data).group(1)
            except:pass                
            try:    
                fulladdress = re.search(r'"displayAddress":"(.*?)"', data).group(1)
            except:pass                
            try:    
                otherurl = re.search(r'"website":(null|".*?")', data).group(1)
            except:pass

            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '4 - paypal', fullname, url, email, user, phone, '', '', fulladdress
                , city, '', '', country, note, '', '', '', '', '', lastname, firstname, '', '', otherurl, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '')        

def pinterest():    # testuser=    kevinrose     # add city
    print(f'{color_yellow}\n\t<<<<< pinterest {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (country, email, fullname,lastname,firstname) = ('', '', '','','')
        (success, note, photo, website, city, otherurls) = ('','','','','', '')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '','','')
        url = ('https://www.pinterest.com/%s/' %(user))
        otherurls = ('https://pinterest.com/search/users/?q=%s' %(user))

        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
           
            fullname = titleurl
            fullname = fullname.split(' (')[0]

            if ' ' in FullName:
                fullname2 = fullname.split(" ")
                firstName = fullname2[0]
                lastName = fullname2[1]
            # success = 1    
        except:
            pass

        if 'Success' in pagestatus:
            for eachline in content.split("\n"):
                if eachline == "": pass                                             # skip blank lines
                else:
                    if 'pinterestapp:about' in eachline and 'rebuildStoreOnClient' not in eachline:
                        eachline = eachline.split('\"')
                        note = eachline[1]

            if note != '':
                write_ossint(user, '4 - pinterest', fullname, url, '', user, '', '', email, ''
                    , city, '', '', country, note, '', '', '', '', '', lastname, firstname, '', '', otherurls, '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)            
                print(f'{color_green} {url}{color_yellow}	   {fullname}	{note}{color_reset}')
                
                
def plaxoemail():    # testEmail= craig@craigslist.org#
    print(f'{color_yellow}\n\t<<<<< plaxo {color_blue}emails{color_yellow} >>>>>{color_reset}')
    for email in emails:
        url = ('http://www.plaxo.com/signup?t=ajax&avail=true&email=%s' %(email))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        if ('Claimed') in content: 
        # if ('Claimed') in str(response): 
            write_ossint(email, '7 - plaxo.com (email exists)', '', url, '', '', '', '', email, ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '')
            print(f'{color_green} {Email}   {Url}{color_reset}')

            
def print_logo():
    
    art = """
 ___    _            _   _ _         _   _             _   
|_ _|__| | ___ _ __ | |_(_) |_ _   _| | | |_   _ _ __ | |_ 
 | |/ _` |/ _ \ '_ \| __| | __| | | | |_| | | | | '_ \| __|
 | | (_| |  __/ | | | |_| | |_| |_| |  _  | |_| | | | | |_ 
|___\__,_|\___|_| |_|\__|_|\__|\__, |_| |_|\__,_|_| |_|\__|
                               |___/                       

  """
    print(f'{color_blue}{art}{color_reset}')


def public():    # testuser=    kevinrose
    print(f'{color_yellow}\n\t<<<<< public {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (country, email, fullname,lastname,firstname) = ('', '', '','','')
        (success, note, photo, website, city, otherurls) = ('','','','','', '')
        (content) = ('')
        url = ('https://public.com/@%s' %(user))
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except:
            pass

        if 'Success' in pagestatus:
            for eachline in content.split("\n"):
                if eachline == "": pass                                             # skip blank lines
                elif "og:title" in eachline:
                    fullname = eachline.strip().split("\"")[1]
                    fullname = fullname.split(" (")[0]
                    if ' ' in fullname:

                        fullname2 = fullname.split(" ")
                        firstname = fullname2[0]
                        lastname = fullname2[1]
              
            if note != 'true':
                write_ossint(user, '4 - public', fullname, url, '', user, '', '', email, ''
                    , city, '', '', country, note, '', '', '', '', '', lastname, firstname, '', '', otherurls, '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, '')            
                print(f'{color_green} {url}{color_yellow}	   {fullname}	{note}{color_reset}')
 
def poshmark():    # testuser=    kevinrose
    print(f'{color_yellow}\n\t<<<<< poshmark {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (country, email, fullname,lastname,firstname) = ('', '', '','','')
        (success, note, photo, website, city, otherurls) = ('','','','','', '')
        (content) = ('')
        url = ('https://poshmark.com/closet/%s' %(user))
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except:
            pass

        if 'Success' in pagestatus:
            for eachline in content.split("\n"):
                if eachline == "": pass                                             # skip blank lines
                elif "og:title" in eachline:
                    fullname = eachline.strip().split("\"")[1].replace('\'s Closet', '')
                    firstname = fullname

            if note != 'true':
                write_ossint(user, '7 - poshmark', fullname, url, '', user, '', '', email, ''
                    , city, '', '', country, note, '', '', '', '', '', lastname, firstname, '', '', otherurls, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '')            
                print(f'{color_green} {url}{color_yellow}	   {fullname}	{note}{color_reset}')
                
def read_url():
    global row  # The magic to pass row globally
    style = workbook.add_format()
    color = 'white'

    inputfile = open(filename)
    # Create dnsdomain
    for eachline in inputfile:
        
        
        (query, ranking, fullname, url, email , user) = ('', '', '', '', '', '')
        (phone, ip, entity, fulladdress, city, state) = ('', '', '', '', '', '')
        (zip, country, note, aka, dob, gender) = ('', '', '', '', '', '')
        (info, misc, lastname, firstname, middlename, friend) = ('', '', '', '', '', '')
        (otherurls, otherphones, otheremails, case, sosfilenumber, president) = ('', '', '', '', '', '')
        (sosagent, managers, dnsdomain, dstip, srcip  ) = ('', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        
        (color) = 'white'
        style.set_bg_color('white')  # test

        eachline = eachline + "," * 72
        eachline = eachline.split(',')  # splits by comas
        note = (eachline[0].strip())

        # Regex data type
        if bool(re.search(r"^[\w\.\+\-]+\@[\w]+\.[a-z]{2,3}$", note)):  # regex email    
            logsource = 'IOC-Email'

        elif re.search(regex_host, note):  # regex_host (Finds url and dnsdomain) # False positives for emails    # todo
            url = note
            if url.lower().startswith('http'):
                logsource = 'IOC-url'
            else:
                logsource = 'IOC-dnsdomain'
            url = url.rstrip('/')
            dnsdomain = url.lower()
            dnsdomain = dnsdomain.replace("https://", "")
            dnsdomain = dnsdomain.replace("http://", "")
            dnsdomain = dnsdomain.split('/')[0]
            notes2 = dnsdomain.split('.')[-1]

        elif re.search(regex_ipv4, note):  # regex_ipv4
            (srcip, dstip) = (note, note)
            logsource = 'IOC-IP'
        elif re.search(regex_md5, note):  # regex_md5
            logsource = 'IOC-Hash-Md5'
        elif re.search(regex_sha256, note):  # regex_sha256
            logsource = 'IOC-Hash-Sha256'
        elif re.search(regex_sha1, note):  # regex_sha1
            logsource = 'IOC-Hash-Sha1'
        elif re.search(regex_sha512, note):  # regex_sha512
            logsource = 'IOC-Hash-Sha512'

        # Read url,dnsdomain or dstip
        if url != '':
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        elif dnsdomain != '':
            (content, referer, osurl, titleurl, pagestatus) = request(dnsdomain)
        elif dstip != '':
            (content, referer, osurl, titleurl, pagestatus) = request(dstip)
        else:
            format_function(bg_color='white')

        # Write OSSINT excel
        write_ossint(query, ranking, fullname, url, email , user, phone, ip, entity, 
            fulladdress, city, state, zip, country, note, aka, dob, gender, info, 
            misc, lastname, firstname, middlename, friend, otherurls, otherphones, 
            otheremails, case, sosfilenumber, president, sosagent, managers, 
            dnsdomain, dstip, srcip, content, referer, osurl,
            titleurl, pagestatus)

        print(f'{color_blue}{note}    {color_green}{osurl}    {color_red}    {titleurl}{color_yellow}    {pagestatus}')


def redirect_detect():
    print(f'{color_yellow}\n\t<<<<< redirected {color_blue}websites{color_yellow} >>>>>{color_reset}')
    for website in websites:    
        (final_url, dnsdomain, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '', '')
        url = website
        url = url.replace("http://", "https://")
        if "http" not in url.lower():
            url = ('https://%s' %(url))
            
        referer = url.lower().strip()
        try:
            response = requests.get(url)

            final_url = response.url

            # print(url, " redirects to ", final_url)        

        except TypeError as error:
            # print(f'{color_red}{error}{color_reset}')
            pass
        
        dnsdomain = url.lower()
        dnsdomain = dnsdomain.replace("https://", "")
        dnsdomain = dnsdomain.replace("http://", "")
        dnsdomain = dnsdomain.split('/')[0]

        ip = ip_address(dnsdomain)
        
        if dnsdomain not in final_url:
            print(f'{color_green}{url} redirects to {final_url}{color_reset}') 
            write_ossint(url, '7 - redirect ', '', final_url, '', '', '', ip, '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', dnsdomain, '', '', '', referer, osurl, titleurl, pagestatus)


def request(url):
    (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
    (string) = ('')
    fake_referer = 'https://www.google.com/'
    headers = {'Referer': fake_referer}
    url = url.replace("http://", "https://")     # test
    # if "http" not in url.lower():
        # url = ('https://%s' %(url))
            
    # if url.lower().startswith('http'):
        # page = requests.get(url)
        
        
    # else:
        # page  = requests.get("http://" +url)

    page  = requests.get(url, headers=headers)
    pagestatus  = page.status_code
    soup = BeautifulSoup(page.content, 'html.parser')
    content = soup.prettify()
    try:
        osurl = page.headers['Server']
    except:pass
    
    try:
        titleurl = soup.title.string
    except:pass
    
    # try:    
        # page  = requests.get(url, headers=headers)
        # pagestatus  = page.status_code
        # soup = BeautifulSoup(page.content, 'html.parser')
        # content = soup.prettify()
        # osurl = page.headers['Server']
        # titleurl = soup.title.string
    # except:
        # pass
    
#pagestatus
    
    if str(pagestatus).startswith('2') :    
        pagestatus = ('Success - %s' %(pagestatus))
    elif str(pagestatus).startswith('3') :    
        pagestatus = ('Redirect - %s' %(pagestatus))
    elif str(pagestatus).startswith('4') :    
        pagestatus = ('Fail - %s' %(pagestatus))
    elif str(pagestatus).startswith('5') :    
        pagestatus = ('Fail - %s' %(pagestatus))
    elif str(pagestatus).startswith('1') :    
        pagestatus = ('Info - %s' %(pagestatus))
    try:
        pagestatus = pagestatus.strip()
    except Exception as e:
        print(f"{color_red}Error striping pagestatus: {str(e)}{color_reset}")
# titleurl

    try:
        title_tag = content.find('title')
        if title_tag is not None:
            title = title_tag.text.strip()
            # The full name is often included in parentheses after the username
            # Example: "Tom Anderson (myspacetom)"
            # We'll split the title at the first occurrence of '(' to extract the full name
            parts = title.split(' (', 1)
            if len(parts) > 1:
                titleurl = parts[0]
    except Exception as e:
        # print(f"Error parsing title: {str(e)}")
        pass

    if titleurl !="":
        try:
            meta_tags = content.find_all('meta')
            for tag in meta_tags:
   
                if tag.get('property') == 'og:title':
                    titleurl = tag.get('content')
                    titleurl =  title.split(' (')[0]
        except Exception as e:
            # print(' ')
            # print(f'{color_red}Error parsing metadata: {str(e)}{color_reset}')
            pass

    try:
        titleurl = str(titleurl)    #test
        titleurl = (titleurl.encode('utf8'))    # 'ascii' codec can't decode byte
        titleurl = (titleurl.decode('utf8'))    # get rid of bytes b''
    except TypeError as error:
        print(f'{color_red}{error}{color_reset}')
        titleurl = ''

    titleurl = titleurl.strip()
    content = content.strip()
    
    return (content, referer, osurl, titleurl, pagestatus)    

def request_url(url):
    
    fake_referer = 'https://www.google.com/'
    headers = {'Referer': fake_referer}

    
    (content, referer, osurl, titleurl, pagestatus)= ('blank', '', '', '', '')
    (response) = ('')
    
    if url.lower().startswith('http'):
        blah = ''
    else:
        url = ("https://" +url)

    try:
        response = requests.get(url, verify=False, headers=headers)        
        response.raise_for_status()
        pagestatus  = response.status_code
        content = response.content.decode()
        content = BeautifulSoup(content, 'html.parser')
        


    except requests.exceptions.RequestException as e:
        # print(f"Could not fetch URL {url}: {str(e)}")
        # raise requests.exceptions.RequestException(str(e))
        (pagestatus) = ('Fail')

# osurl
    try:
        osurl = response.headers['Server']
    except:
    # except KeyError:
        pass

# titleurl
    try:
        titleurl = titleurl_get(content)
        # titleurl = titleurl_og(content)
    except KeyError:
        titleurl = ''


#pagestatus
    
    if str(pagestatus).startswith('2') :    
        pagestatus = ('Success - %s' %(pagestatus))
    elif str(pagestatus).startswith('3') :    
        pagestatus = ('Redirect - %s' %(pagestatus))
    elif str(pagestatus).startswith('4') :    
        pagestatus = ('Fail - %s' %(pagestatus))
    elif str(pagestatus).startswith('5') :    
        pagestatus = ('Fail - %s' %(pagestatus))
    elif str(pagestatus).startswith('1') :    
        pagestatus = ('Info - %s' %(pagestatus))

    pagestatus = pagestatus.strip()    

    return (content, referer, osurl, titleurl, pagestatus)

def resolverRS():# testIP= 77.15.67.232
    print(f'{color_yellow}\n\t<<<<< resolverRS {color_blue}ip{color_yellow} >>>>>{color_reset}')
   
    for ip in ips:
        (country, city, zip, case, note, state) = ('', '', '', '', '', '')
        (misc, info) = ('', '')
        
        url = ('https://resolve.rs/ip/geolocation.html?ip=%s' %(ip))
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "403 ERROR" in eachline:
                pagestatus = '403 Error'
                content = ''

            elif "\"code\": \"" in eachline :   # and zip != ''
                zip = eachline.split("\"")[3]
            elif "\"en\": \"" in eachline :   # city
                # print(f'')
                if city == '':
                    city = eachline.split("\"")[3]
                elif misc == '' and city != '': # continent
                    misc = eachline.split("\"")[3]

                elif misc != '' and country == '' and city != '': # country
                    country = eachline.split("\"")[3]
                elif info == '' and misc != '' and country != '' and city != '':    # registered country
                    info = eachline.split("\"")[3]
                elif state == '' and info != '' and misc != '' and country != '' and city != '':    # state
                    state = eachline.split("\"")[3]
            elif "COMCAST" in eachline :   # isp  >ASN</a>
                note = 'COMCAST'


        # pagestatus = ''                
        if url != '':
            print(f'{color_green}{url}{color_reset}') 
            write_ossint(ip, '6 - resolve.rs', '', url, '', '', '', ip, '', ''
                , city, state, zip, country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, pagestatus)

def reversephonecheck():# testPhone= 708-372-8101   https://www.reversephonecheck.com/1-708/372/81/#01
    print(f'{color_yellow}\n\t<<<<< reversephonecheck {color_blue}phone numbers{color_yellow} >>>>>{color_reset}')
    
    for phone in phones:
        (query) = (phone)
        (fulladdress, country, city, case, note) = ('', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '')
        (areacode, prefix, line, count, match, match2) = ('', '', '', 1, '', '')
        (url) = ('')
        phone = phone.replace('(','').replace(')','-').replace(' ','')
        if phone.startswith('1-'):
            phone = phone.replace('1-','')
        elif phone.startswith('1'):
            phone = phone.lstrip('1')

        if len(phone) != 10:
            # print(f'{color_red}Invalid phone number{color_reset} {phone}')
            print('')
        elif '-' not in phone:
            phone = (phone[:3] + "-" + phone[3:6] + "-" + phone[6:])
  
        (line2) = ('')
        if "-" in phone:
            phone2 = phone.split("-")
            areacode = phone2[0]
            prefix = phone2[1]
            try:
                line = phone2[2]
                line2 = line
                line = line[:2]
                
            except:
                pass
        url = ('https://www.reversephonecheck.com/1-%s/%s/%s/#%s' %(areacode, prefix, line, phone.replace('-', ''))) 
        # print(f'{color_yellow}url =  {url}{color_reset}')   #temp
        # print(f'{color_blue}{phone}{color_reset}') # temp
        (content, referer, osurl, titleurl, pagestatus) = request(url) 
        match = ("%s - %s" %(prefix, line2))

        for eachline in content.split("\n"):
            if match in eachline:
                pagestatus = 'research'
                count += 1

        if pagestatus == 'research' and count == 2:
            print(f'{color_green}{url}{color_reset}')
            write_ossint(query, '3 - reversephonecheck', '', url, '', '', phone, '', '', fulladdress
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)

def robtex():
    print(f'{color_yellow}\n\t<<<<<robtex dns lookup >>>>>{color_reset}')    

    for website in websites:    
        (final_url, dnsdomain, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '', '')
        (otherurl, ip) = ('', '')
        url = website
        url = url.replace("http://", "https://")
        if "http" not in url.lower():
            url = ('https://%s' %(url))
        
        dnsdomain = url.lower()
        dnsdomain = dnsdomain.replace("https://", "")
        dnsdomain = dnsdomain.replace("http://", "")
        dnsdomain = dnsdomain.split('/')[0]

        ip = ip_address(dnsdomain)

        otherurl = url  

        url = ('https://www.robtex.com/dns-lookup/%s#quick' %(dnsdomain))
        
        if 1==1:
        # if dnsdomain not in final_url:
            print(f'{color_green}{website}{color_yellow}	{ip}{color_reset}')
            write_ossint(otherurl, '9 - robtexDNS-lookup ', '', url, '', '', '', ip, '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', otherurl, '', '', '', '', '', '', '', dnsdomain, '', '', '', '', '', '', '')
   

def samples():
    print(f'''{color_yellow}    
Alain_THIERRY
anh_usa
frizz925
gurkha_love
Hmei7
itdoesnthavetohappen
JLight29
kevinrose
kevinwagner
kobegal
luckyjames844
maverick3819
merz356
Misty0427
MR-JEFF
N3tAtt4ck3r
Pattycakes98
rothys
ryanlwatkins
thekevinrose
williger
zazenergy
nullcrew
realDonaldTrump
Their1sn0freakingwaythisisreal12JT4321
        
77.15.67.232
92.20.236.78
255.255.255.255

annemconnor@yahoo.com
kandyem@yahoo.com
kevinrose@gmail.com
craig@craigslist.org
ceo@zappos.com
lnd_whitaker@yahoo.com
gsmstocks@gmail.com
lydianorman1@hotmail.com
soniraj388@gmail.com
tanderson09@gmail.com
tin_max87@yahoo.com
Their1sn0freakingwaythisisreal12344321@fakedomain.com

385-347-1531
999-999-9999
7083703020
{color_reset}
'''
)    

def sherlock():    # testuser=    kevinrose
    print(f'\n\t{color_yellow}<<<<< Manually check Sherlock users >>>>>{color_reset}')
    
    for user in users:    
        note = ('cd C:\Forensics\scripts\python\git-repo\sherlock && python sherlock %s' %(user)) 

        if 1==1:

            write_ossint(user, '8 - manual', '', '', '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', 'research')


def snapchat(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< snapchat {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', 'research', '')
        user = user.rstrip()
        url = ('https://www.snapchat.com/add/%s?' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):

            if "og:title" in eachline:
                fullname = eachline.strip().split("\"")[1].replace(' on Snapchat','')


        if 'name=\"description' in content:
            write_ossint(user, '6 - snapchat', fullname, url, '', user, '', '', '', ''
            , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)        
            print(f'{color_green}{url}{color_reset}')        

def spotify(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< spotify {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = ('https://open.spotify.com/user/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        if '404' not in pagestatus:
            titleUrl = titleurl.replace(" on Spotify","").strip()
            fullname = titleurl
            print(f'{color_green}{url}{color_yellow}	{titleurl}{color_reset}') 
            write_ossint(user, '6 - spotify', '', url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)        

def spydialer():# testPhone= 708-372-8101
    print(f'{color_yellow}\n\t<<<<< spydialer {color_blue}users{color_yellow} >>>>>{color_reset}')

    for phone in phones:
        (country, city, zip, case, note) = ('', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '')
        phone = phone.replace('(','').replace(')','').replace(' ','')

        url = ('https://www.spydialer.com')

        pagestatus = 'research'  
        print(f'{color_yellow}{phone}{color_reset}')
        write_ossint(phone, '3 - spydialer', '', url, '', '', phone, '', '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)

def thatsthememail():# testEmail= smooth8101@yahoo.com   
    print(f'{color_yellow}\n\t<<<<< thatsthem {color_blue}emails{color_yellow} >>>>>{color_reset}')
    
    for email in emails:
        print(f'{color_red}{email}{color_reset}')
        (country, city, zip, case, note) = ('', '', '', '', '')
        
        url = ('https://thatsthem.com/email/%s' %(email))
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "Found 0 results for your query" in eachline and case == '':
                # print(f'{color_red}not found{color_reset}')  # temp
                url = ('')
                
        pagestatus = ''                
        if url != '':
            print(f'{color_green}{url}{color_yellow}	{email}{color_reset}') 
            
            
            write_ossint(email, '9 - thatsthem', '', url, email, '', '', '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, pagestatus)
        else:
            print(f'{color_red}{url}{color_yellow}	{email}{color_reset}') 
            
            
            write_ossint(email, '99 - thatsthem', '', url, email, '', '', '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, pagestatus)
        
def thatsthemip():# testIP= 8.8.8.8
    print(f'{color_yellow}\n\t<<<<< thatsthem {color_blue}ip{color_yellow} >>>>>{color_reset}')
       
    for ip in ips:
        (country, city, zip, case, note, state) = ('', '', '', '', '', '')
        
        url = ('https://thatsthem.com/ip/%s' %(ip))
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "located in " in eachline:
                state = eachline
                note = eachline

            elif "403 ERROR" in eachline:
                pagestatus = '403 Error'
                content = ''
            elif "Found 0 results for your query" in eachline:
                print(f'{color_red}Not found{color_reset}')  # temp
                url = ('')
        # pagestatus = ''                
        if url != '':
            print(f'{color_green}{url}{color_reset}') 
            write_ossint(ip, '6 - thatsthem', '', url, '', '', '', ip, '', ''
                , city, state, '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, pagestatus)


def thatsthemphone():# testPhone= 708-372-8101  
    print(f'{color_yellow}\n\t<<<<< thatsthem {color_blue}phone numbers{color_yellow} >>>>>{color_reset}')
    for phone in phones:
        (country, city, zip, case, note) = ('', '', '', '', '')
        phone = phone.replace('(','').replace(')','-')
        time.sleep(10) # will sleep for 10 seconds
        url = ('https://thatsthem.com/phone/%s' %(phone))    # https://thatsthem.com/reverse-phone-lookup
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        if "Found 0 results for your query" in content or "The request could not be satisfied" in content:
            # print(f'{color_red}{url}{color_yellow}	Found 0 results for your query{color_reset}')   # temp
            # url = ('')
            note = ('captcha protected')

        for eachline in content.split("\n"):
            if "Found 0 results for your query" in eachline and case == '':
                print(f'{color_red}Not found{color_reset}')  # temp
                # url = ('')
        pagestatus = ''        
                
        if note == '':
            print(f'{color_green}{url}{color_reset}')
            write_ossint(phone, '6 - thatsthem', '', url, '', '', phone, '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)
        else:   
            print(f'{color_yellow}{url}{color_reset}')
            write_ossint(phone, '9 - thatsthem', '', url, '', '', phone, '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)


def telegram(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< telegram {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', 'research', '')
        user = user.rstrip()
        url = ('https://t.me/%s' %(user))
        
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            # (content, referer, osurl, titleurl, pagestatus) = request_url(url)

            for eachline in content.split("\n"):
                if "og:title" in eachline:
                    fullname = eachline.strip().split("\"")[1]

            if 'Telegram' not in fullname:
                print(f'{color_green}{url}{color_yellow}	{titleurl}{color_reset}') 
                write_ossint(user, '7 - telegram', fullname, url, '', user, '', '', '', ''
                    , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        
        except TypeError as error:
            print(f'{color_red}{error}{color_reset}')
            
def threads():    # testuser=    kevinrose     # add info
    print(f'{color_yellow}\n\t<<<<< threads {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        url = ('https://www.threads.net/@%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
 
            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    fullname = eachline.split('\"')[1].split(' (')[0]
                elif "og:description" in eachline:
                    note = eachline.strip()
                    note = note.replace("\" property=\"og:description\"/>",'').replace("<meta content=\"",'')

        except:
            pass
        if 'on Threads' in titleurl:

            write_ossint(user, '3 - threads', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_reset}')    

def tiktok(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< tiktok {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', 'research', '')
        user = user.rstrip()
        url = ('https://tiktok.com/@%s?' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        if 'uccess' in pagestatus:
            fullname = titleurl
            fullname = fullname.split(' (')[0]
            if fullname == user:
                fullname = ''
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '4 - tiktok', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        

def tinder():    # testuser=    john
    print(f'{color_yellow}\n\t<<<<< tinder {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        url = ('https://tinder.com/@%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, dob, info, misc, note) = ('', '', '', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            
            
            
            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    fullname = eachline.split('\"')[1].split(' (')[0]
                elif 'schools\"' in eachline:
                # elif 'schools\":\[{\"name' in eachline:
                    data = json.loads(eachline)
                    note = data["schools"][0]["name"]
                    # Extract the description value and print it
                    # note = data['description']
                    misc = eachline
                    
                    dob= data["webProfile"]["user"]["birth_date"]
        except:
            pass
            
        # time.sleep(1) # will sleep for 1 seconds
        if 'alternate' in content:
            
        
            write_ossint(user, '7 - tinder.com', fullname, url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', dob, '', info, misc, '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_reset}')    



def titles():    # testsite= google.com
    from subprocess import call, Popen, PIPE
    print(f'{color_yellow}\n\t<<<<< Titles grab {color_blue}Website\'s{color_yellow} >>>>>{color_reset}')
    for website in websites:    
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        (ip, note) = ('', '')
        
        url = website

        fake_referer = 'https://www.google.com/'
        headers = {'Referer': fake_referer}
        url = url.replace("http://", "https://")     # test
        if "http" not in url.lower():
            url = ('https://%s' %(url))

        url = url.replace("https://", "http://")

        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except TypeError as error:
            print(f'{color_red}{error}{color_reset}')
        
        # dnsdomain
        dnsdomain = url.lower()
        dnsdomain = dnsdomain.replace("https://", "")
        dnsdomain = dnsdomain.replace("http://", "")
        dnsdomain = dnsdomain.split('/')[0]
        
        # ip
        ip = ip_address(dnsdomain)
        print(f'{color_green}{website}{color_yellow}	   {pagestatus}	{color_blue}{titleurl}{color_reset}')

        write_ossint(url, '7 - website ', '', url, '', '', '', ip, '', ''
            , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', dnsdomain, '', '', '', referer, osurl, titleurl, pagestatus)


def titleurl_get(content):

    titleurl = ''
    try:
        # soup = BeautifulSoup(html, 'html.parser')
        title_tag = content.find('title')
        if title_tag is not None:
            title = title_tag.text.strip()
            # The full name is often included in parentheses after the username
            # Example: "Tom Anderson (myspacetom)"
            # We'll split the title at the first occurrence of '(' to extract the full name
            parts = title.split(' (', 1)
            if len(parts) > 1:
                titleurl = parts[0]
    except Exception as e:
        # print(f'Error parsing title: {str(e)}')
        pass


    return titleurl

def titleurl_og(content):
    (titleurl) = ('')

    try:
        meta_tags = content.find_all('meta')
        for tag in meta_tags:
            if tag.get('property') == 'og:title':
                titleurl = tag.get('content')
                titleurl =  title.split(' (')[0]
    except Exception as e:
        pass
    return titleurl

def truthSocial(): # testuser = realdonaldtrump https://truthsocial.com/@realDonaldTrump
    print(f'{color_yellow}\n\t<<<<< truthsocial {color_blue}users{color_yellow} >>>>>{color_reset}')
    print(f'{color_yellow}\n\t\t\tThis one one takes a while{color_reset}')
    for user in users:    
        (city, country, note, fullname, titleurl, pagestatus) = ('', '', '', '', '', '')
        (info, ranking) = ('', '9 - truthsocial.com')
        user = user.rstrip()
        url = ('https://truthsocial.com/@%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        pagestatus = ''
        time.sleep(3) #will sleep for 3 seconds
        for eachline in content.split("  <"):
            if 'This resource could not be found' in eachline:
                pagestatus = '404'
            elif "og:title" in eachline:
                titleurl = eachline.strip().split("\"")[1]
                fullname = titleurl.split(" (")[0]
                pagestatus = '200'
                ranking = '9 - truthsocial'
                if titleurl == 'Truth Social':
                    pagestatus = '404'
                else:
                    pagestatus = '200'
                    ranking = '3 - truthsocial'
            elif "og:description" in eachline:
                note = eachline.strip().split("\"")[1]

        if '@' in titleurl: 
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, ranking, fullname, url, '', user, '', '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)  

        
def twitter():    # testuser=    kevinrose     # add info
    print(f'{color_yellow}\n\t<<<<< twitter {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (fullname,lastname,firstname, email, city, country) = ('','','', '', '', '')
        url = ('https://twitter.com/%s' %(user))
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        # print(titleurl, url, pagestatus)  # temp
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            # print(titleurl)  # temp
            titleurl = titleurl.replace(") on Twitter","")
            titleurl = titleurl.lower().replace(User.lower(),"")
            titleurl = titleurl.replace(" (","")
            fullname = titleurl
            fullname = fullname.replace(" account suspended","")
            fullname = fullname.replace("twitter /","")
            titleurl = titleurl.lower().replace(fullname.lower(),"")

            write_ossint(user, '5 - twitter', fullname, url, email, user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_green}{url}{color_yellow}	   {fullname}	{titleurl}{color_reset}')
        except:
            write_ossint(user, '9 - twitter.com', fullname, url, email, user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)
            print(f'{color_yellow}{url}{color_yellow}	   {fullname}	{titleurl}{color_reset}')

        time.sleep(10) #will sleep for 10 seconds
            
def twitteremail(): # test Email=     craig@craigslist.org 
    print(f'{color_yellow}\n\t<<<<< twitter {color_blue}emails{color_yellow} >>>>>{color_reset}')

    # {"valid":false,"msg":"Email has already been taken. An email can only be used on one Twitter account at a time.","color":"red","taken":true,"blank":false}
    for email in emails:
        url = ('https://twitter.com/users/email_available?email=%s' %(email))
        (country, city, zip) = ('', '', '')        
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)    # access denied cloudflare
        except socket.error as ex:
            print(ex)        
        if 'Email has already been taken' in content:
            write_ossint(email, '5 - twitter.com', '', url, email, '', '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '')
            print(f'{color_green}{email}{color_reset}')

def validnumber():# testPhone= 7083703020
    print(f'{color_yellow}\n\t<<<<< validnumber {color_blue}phone numbers{color_yellow} >>>>>{color_reset}')

    # https://validnumber.com/phone-number/3124377966/
    for phone in phones:
        (country, city, state, zip, case, note) = ('', '', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '')
        (query) = (phone)
        phone = phone.replace('(','').replace(')','').replace(' ','')
        if phone.startswith('1-'):
            phone =phone.replace('1-','')
        elif phone.startswith('1'):
            phone =phone.lstrip('1')
 
        url = ('https://validnumber.com/phone-number/%s/' %(phone.replace("-", "")))
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)    # protected by cloudflare
        pagestatus = ''
        for eachline in content.split("\n"):
            if "No name associated with this number" in eachline and case == '':
                print(f'{color_red}not found{color_reset}')  # temp
                url = ('')
            elif "Find out who owns" in eachline:
                if 'This device is registered in ' in eachline:
                    note = eachline.split('\"')[1]
                    note = note.split('Free owner details for')[0]
                    city = eachline.split("This device is registered in ")[1].split("Free owner details")[0]
                    
                    state = city.split(',')[1]
                    city = city.split(',')[0]

        if city != '':        
            print(f'{color_green}{url}{color_reset}') 
            write_ossint(query, '5 - validnumber', '', url, '', '', phone, '', '', ''
                    , city, state, '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, '')

def viewdnsdomain():
    print(f'{color_yellow}\n\t<<<<<viewdns lookup >>>>>{color_reset}')    

    for website in websites:    
        (final_url, dnsdomain, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        (otherurl, ip) = ('', '')
        url = website
        url = url.replace("http://", "https://")
        if "http" not in url.lower():
            url = ('https://%s' %(url))
        
        dnsdomain = url.lower()
        dnsdomain = dnsdomain.replace("https://", "")
        dnsdomain = dnsdomain.replace("http://", "")
        dnsdomain = dnsdomain.split('/')[0]

        ip = ip_address(dnsdomain)

        otherurl = url  

        url = ('https://viewdns.info/whois/?domain=%s' %(dnsdomain))
        # (content, referer, osurl, titleurl, pagestatus) = request(url)
        # time.sleep(10) #will sleep for 10 seconds
        if 1==1:
        # if dnsdomain not in final_url:
            print(f'{color_green}{website}{color_yellow}	{ip}{color_reset}')
            write_ossint(url, '9 - viewdns ', '', url, '', '', '', ip, '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', otherurl, '', '', '', '', '', '', '', dnsdomain, '', '', '', '', '', titleurl, pagestatus)
   


def whatismyip():    # testuser= 77.15.67.232  
    print(f"{color_yellow}\n\t<<<<< whatismyipaddress.com {color_blue}IP's{color_yellow} >>>>>{color_reset}")
    for ip in ips:
        (country, city, state, zip, pagestatus, title) = ('', '', '', '', '', '')
        url = ('https://whatismyipaddress.com/ip/%s' %(ip))
        (content, titleurl) = ('', '')

        # time.sleep(7) #will sleep for 30 seconds
        print(f'{color_green}{ip}{color_reset}')
        write_ossint(ip, '9 - whatismyipaddress', '', url, '', '', '', ip, '', ''
        , city, state, zip, country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)



def whatsmyname():    # testuser=    kevinrose
    print(f'\n\t{color_yellow}<<<<< Manually check whatsmyname users >>>>>{color_reset}')
    for user in users:    
        url = ('https://whatsmyname.app/')
        
        note = ('cd C:\Forensics\scripts\python\git-repo\WhatsMyName && python web_accounts_list_checker.py -u %s -of C:\Forensics\scripts\python\output_%s.txt' %(user, user)) 
            
        # time.sleep(1) # will sleep for 1 seconds
        if 1==1:

            write_ossint(user, '9 - manual', '', url, '', user, '', '', '', ''
                , '', '', '', '', note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', 'research')


def whitepagesphone():# testuser=    210-316-9435
    print(f'{color_yellow}\n\t<<<<< whitepages {color_blue}phone numbers{color_yellow} >>>>>{color_reset}')
    for phone in phones:    
        (country, city, zip) = ('', '', '')
        (titleurl) = ('')
        # url = ('http://www.whitepages.com/search/ReversePhone?full_phone=%s' %(phone))
        url = ('https://www.whitepages.com/phone/1-%s' %(phone))

        # (content, referer, osurl, titleurl, pagestatus) = request(url)    # access denied cloudflare
        print(f'{color_yellow}{phone}{color_reset}')
        write_ossint(phone, '9 - whitepages', '', url, '', '', phone, '', '', ''
            , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, '')
    
def whocalld():# testPhone= 708-372-8101 DROP THE LEADING 1
    print(f'{color_yellow}\n\t<<<<< whocalld {color_blue}phone numbers{color_yellow} >>>>>{color_reset}')

    # https://whocalld.com/+17083728101
    for phone in phones:
        (country, city, state, zip, case, note) = ('', '', '', '', '', '')
        (fullname, content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '', '')
        phone = phone.replace('(','').replace(')','').replace(' ','')
        
        if phone.startswith('1'):
            phone = phone.replace('1','')
 
        url = ('https://whocalld.com/+1%s' %(phone.replace("-", "")))
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)    # protected by cloudflare

        for eachline in content.split("\n"):
            if "Not found" in eachline and case == '':

                url = ('')
            elif "This seems to be" in eachline:
                if ' in ' in eachline:
                    note = eachline.replace(". </p>",'').replace("<p>",'').strip().replace("This",phone)
                    city = eachline.split(" in ")[1].replace(". </p>",'').replace("<p>",'').strip()
                    if ", " in city:
                        state = city.split(", ")[1].replace(".",'')
                        city = city.split(", ")[0]
                    note = ("According to %s %s" %(url, note))
            elif "The name of this caller seemed to be " in eachline:
                note = eachline
                fullname = eachline.replace("The name of this caller seemed to be ",'').split(",")[0].strip()
                if ' in ' in eachline:
                    note = eachline.replace(". </p>",'').replace("<p>",'').strip().replace("This",phone)
                    city = eachline.split(" in ")[2].replace(". </p>",'').replace("<p>",'').strip()
                    if ", " in city:
                        state = city.split(", ")[1].replace(".",'')
                        city = city.split(", ")[0]

        pagestatus = ''        
                
        if url != '':
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(phone, '4 - whocalld', fullname, url, '', '', phone, '', '', ''
                , city, state, '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', referer, '', titleurl, pagestatus)

def whoisip():    # testuser=    77.15.67.232   only gets 403 Forbidden
    from subprocess import call, Popen, PIPE
    print(f"{color_yellow}\n\t<<<<< whois {color_blue}IP's{color_yellow} >>>>>{color_reset}")
    for ip in ips:    
        (city, business, country, zip, state) = ('', '', '', '', '')
        (content, titleurl, pagestatus) = ('', '', '')
        (email, phone, fullname, entity, fulladdress) = ('', '', '', '', '') 
        url = ('https://www.ip-adress.com/whois/%s' %(ip))
        if sys.platform == 'win32' or sys.platform == 'win64':    
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            if '403 Forbidden' in content:
                pagestatus = '403 Forbidden'
                content = ''
            for eachline in content.split("\n"):
                
                if "www.ip-adress.com/legal-notice" in eachline:
                    for eachline in content.split("<tr><th>"):
                        if "Country<td>" in eachline:
                            country = eachline.strip().split("Country<td>")[1]
                        elif "City<td>" in eachline:
                            city = eachline.strip().split("City<td>")[1]
                        elif "ISP</abbr><td>" in eachline:
                            business = eachline.strip().split("ISP</abbr><td>")[1]
                        elif "Postal Code<td>" in eachline:
                            zip = eachline.strip().split("Postal Code<td>")[1]
                        elif "State<td>" in eachline:
                            state = eachline.strip().split("State<td>")[1]
               
                elif "<tr><th>Country</th><td>" in eachline:
                    country = eachline.strip().split("<tr><th>Country</th><td>")[1]
                    country = country.split("<")[0]
                elif "City: " in eachline:
                    city = eachline.strip().split("<tr><th>City</th><td>")[1]
                    city = city.split("<")[0]            
                elif "<tr><th>Postal Code</th><td>" in eachline:
                    zip = eachline.strip().split("<tr><th>Postal Code</th><td>")[1]
                    zip = zip.split("<")[0] 
            time.sleep(3) #will sleep for 30 seconds
            if "Fail" in pagestatus:
                pagestatus = 'fail'
        else:
            WhoisArgs = ('whois %s' %(ip))
            response= Popen(WhoisArgs, shell=True, stdout=PIPE)
            for line in response.stdout:
                line = line.decode("utf-8")
                if ':' in line and "# " not in line and len(line) > 2:
                    line = line.strip()
                    content = ('%s\n%s' %(content, line))
                if email == '':
                    if line.startswith('RAbuseEmail:'):
                        try:
                            email = (line.split(': ')[1].lstrip())
                        except:pass    
                    elif line.lower().startswith('abuse-mailbox:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabuseemail:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('Orgtechemail:'):email = (line.split(': ')[1].lstrip())
                
                if phone == '':
                    if line.lower().startswith('rabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('phone:'):phone = (line.split(': ')[1].lstrip())
                    phone = phone.replace("+", "")
                if line.lower().startswith('rtechname:'):fullname = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('person:'):fullname = (line.split(': ')[1].lstrip())                
                
                if line.lower().startswith('country:'):country = (line.split(': ')[1].lstrip())
                if line.lower().startswith('city:'):city = (line.split(': ')[1].lstrip())
                if line.lower().startswith('address:'):fulladdress = ('%s %s' %(fulladdress, line.split(': ')[1].lstrip()))
                if line.lower().startswith('stateprov:'):state = (line.split(': ')[1].lstrip())
                if line.lower().startswith('postalcode:'):zip = (line.split(': ')[1].lstrip())
                if line.lower().startswith('orgname:'):entity = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('org-name:'):entity = (line.split(': ')[1].lstrip())

        print(f'{color_green}{ip}{color_yellow}	{country}	{city}	{zip}{color_reset}')
        
        write_ossint(ip, '9 - whois', fullname, url, email, '', phone, ip, entity, fulladdress
            , city, state, zip, country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)


def whoiswebsite():    # testsite= google.com
    from subprocess import call, Popen, PIPE
    print(f'{color_yellow}\n\t<<<<< whois {color_blue}Website\'s{color_yellow} >>>>>{color_reset}')

    for dnsdomain in dnsdomains:    
        # query = website
        # website = website.replace('http://','')
        # website = website.replace('https://','')
        # website = website.replace('www.','')
        
        url = ('https://www.ip-adress.com/website/%s' %(dnsdomain))
        url2 = ('https://whois.domaintools.com/%s' %(dnsdomain.replace('www.','')))
        (email,phone,fullname,country,city,state) = ('','','','','','')
        (city, country, zip, state, ip) = ('', '', '', '', '')
        (content, titleurl, pagestatus) = ('', '', '')
        (email, phone, fullname, entity, fulladdress) = ('', '', '', '', '') 

        if sys.platform == 'win32' or sys.platform == 'win64':    
            print(f'skipping whois query from windows')  # temp
            write_ossint(dnsdomain, '7 - whois', fullname, url, email, '', phone, ip, entity, fulladdress
                , city, state, zip, country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', dnsdomain, '', '', '', '', '', titleurl, pagestatus)

            write_ossint(dnsdomain, '9 - whois.domaintools.com', fullname, url2, '', '', '', '', '', ''
                , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', dnsdomain, '', '', '', '', '', '', '')


        elif dnsdomain.endswith('.com') or dnsdomain.endswith('.edu') or dnsdomain.endswith('.net'):
            WhoisArgs = ('whois %s' %(dnsdomain))
            response= Popen(WhoisArgs, shell=True, stdout=PIPE)
            for line in response.stdout:
                line = line.decode("utf-8")
                if ':' in line and "# " not in line and len(line) > 2:
                    line = line.strip()
                    content = ('%s\n%s' %(content, line))
                if email == '':
                    if line.startswith('RAbuseEmail:'):
                        try:
                            email = (line.split(': ')[1].lstrip())
                        except:pass    
                    elif line.lower().startswith('abuse-mailbox:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabuseemail:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('Orgtechemail:'):email = (line.split(': ')[1].lstrip())
                
                if phone == '':
                    if line.lower().startswith('rabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('phone:'):phone = (line.split(': ')[1].lstrip())
                    phone = phone.replace("+", "")
                if line.lower().startswith('rtechname:'):fullname = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('person:'):fullname = (line.split(': ')[1].lstrip())                
                
                if line.lower().startswith('country:'):country = (line.split(': ')[1].lstrip())
                if line.lower().startswith('city:'):city = (line.split(': ')[1].lstrip())
                if line.lower().startswith('address:'):fulladdress = ('%s %s' %(fulladdress, line.split(': ')[1].lstrip()))
                if line.lower().startswith('stateprov:'):state = (line.split(': ')[1].lstrip())
                if line.lower().startswith('postalcode:'):zip = (line.split(': ')[1].lstrip())
                if line.lower().startswith('orgname:'):entity = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('org-name:'):entity = (line.split(': ')[1].lstrip())

            print(f'{color_green}"whois {dnsdomain}{color_yellow}	   {email}	{color_blue}{phone}{color_reset}')
            
            write_ossint(dnsdomain, '7 - whois', fullname, url, email, '', phone, ip, entity, fulladdress
                , city, state, zip, country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', dnsdomain, '', '', '', '', '', titleurl, pagestatus)
        else:
            print(f'{color_red}{dnsdomain} not an edu net or edu site?{color_reset}')
            
            write_ossint(dnsdomain, '7 - whois', fullname, url, email, '', phone, ip, entity, fulladdress
                , city, state, zip, country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', dnsdomain, '', '', '', '', '', titleurl, pagestatus)
        time.sleep(7)

def wordpress(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< wordpress {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (Success, fullname, lastname, firstname, case, gender) = ('','','','','','')
        (photo, country, website, email, language, username) = ('','','','','','')
        (city, note) = ('', '')
        user = user.rstrip()
        url = ('https://wordpress.org/support/users/%s/' %(user))
        note = ('https://%s.wordspress.com' %(user))        
        
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except socket.error as ex:
            print(f'{color_red}{ex}{color_reset}')

        if 'That page can' not in content:
        # if 'Do you want to register' not in content:
            titleUrl = titleurl.replace("'s Profile | WordPress.org","").strip()
            fullname = titleurl
            fullname = fullname.split(" (")[0]
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '4 - wordpress', '', url, '', user, '', '', '', ''
                , city, '', '', country, note, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus)

def wordpressprofiles(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< wordpress {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (Success, fullname, lastname, firstname, case, gender) = ('','','','','','')
        (photo, country, website, email, language, username) = ('','','','','','')
        (city) = ('')
        user = user.rstrip()
        url = ('https://profiles.wordpress.org/%s/' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except:
            pass
        if '404' not in pagestatus:
            fullname = titleurl
            fullname = fullname.split(" (")[0]
            if fullname.lower() in titleurl.lower():
                (fullname, titleurl) = ('', '')
            
            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}') 
            write_ossint(user, '5 - wordpress', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)

                
def wordpresssearchemail():    # testuser=    kevinrose@gmail.com 
    print(f'{color_yellow}\n\t<<<<< wordpressemail {color_blue}emails{color_yellow} >>>>>{color_reset}')    
    
    for email in emails:
        
        url = ('http://en.search.wordpress.com/?q=%s' %(email))

        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except:
            pass

        if 'Your search did not match any blog posts' not in content:

            content = content.split('\n') 
            for eachline in content:

                if eachline == "": pass                                             # skip blank lines
                else:
                    if 'post-title' in eachline:
                        # print(eachline) # temp
                        eachline = eachline.split('\"')
                        url = eachline[3]
                        write_ossint(email, '9 - wordpress', '', url, email, '', '', '', '', ''
                            , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', 'research')
                        print(f'{color_green}{url}{color_reset}')    
        else:   # temp testing
            write_ossint(email, '99 - wordpress', '', url, email, '', '', '', '', ''
                            , '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', pagestatus, 'research')
            print(f'{color_red}{url}{color_reset}')    
            
def write_blurb():
    '''
    read intel.xlsx and write ossint_.docx to describe what you found.
    '''
    print('hello world')     # temp
    excel_file = "intel.xlsx"
    docx_file = "ossint_.docx"

    if not os.path.exists(excel_file):
        input(f"{color_red}{excel_file} doesnt exist.{color_reset}")
        sys.exit()
    elif os.path.getsize(excel_file) == 0:
        input(f'{color_red}{excel_file} is empty. Fill it with username, email, ip, phone and/or websites.{color_reset}')
        sys.exit()
    elif os.path.isfile(excel_file):
        print(f'{color_green}Reading {excel_file}{color_reset}')        
        # write_blurb(excel_file, docx_file)
    else:
        input(f'{color_red}See {excel_file} does not exist. Hit Enter to exit...{color_reset}')
        sys.exit()


    # Open the Excel file
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb.active
    
    # Find the column headers
    header_row = sheet[1]
    column_names = [cell.value for cell in header_row]    
    
    # Columns to skip
    columns_to_skip = ["ranking", "content", "titleurl", "pagestatus"]
    
    for idx, cell in enumerate(header_row, start=1):
        if cell.value == "fullname":
            fullname_column_index = idx
            # print(f'Fullname: {fullname_column_index}')
            break

    if fullname_column_index is None:
        print("Fullname column not found in the Excel file.")
        return

    # Create a new Word document
    doc = Document()
 
    sentence = (f'An open source search revealed the following details.\n\n')
    print(f'{sentence}')  
    doc.add_paragraph(sentence)    
    # Loop through rows in the Excel file and write to Word document
    for row in sheet.iter_rows(min_row=2, values_only=True):
        sentence = "\n".join(f"{column}: {value}" for column, value in zip(column_names, row) if column not in columns_to_skip and value is not None)
        doc.add_paragraph(sentence)
        doc.add_paragraph("")  # Add an empty line between rows


    # Save the Word document
    doc.save(docx_file)
    print(f"Data written to {docx_file}")
    
def write_ossint(query, ranking, fullname, url, email , user, phone, ip, entity, 
    fulladdress, city, state, zip, country, note, aka, dob, gender, info, 
    misc, lastname, firstname, middlename, friend, otherurls, otherphones, 
    otheremails, case, sosfilenumber, president, sosagent, managers, 
    dnsdomain, dstip, srcip, content, referer, osurl,
    titleurl, pagestatus):

    # color
    try:
        if 'Fail' in pagestatus:  #
            format_function(bg_color='red')
        elif 'fail' in pagestatus or 'research' in pagestatus:  
            # format_function(bg_color='orange')
            format_function(bg_color='#FFc000')         # orange  
    except:pass        
        
        
    # elif 'google' in url:  #
        # format_function(bg_color='#92D050')  # green
    else:
        format_function(bg_color='white')
    
    global row

    sheet1.write_string(row, 0, query, format)
    sheet1.write_string(row, 1, ranking, format)
    sheet1.write_string(row, 2, fullname, format)
    sheet1.write_string(row, 3, url, format)
    sheet1.write_string(row, 4, email, format)
    sheet1.write_string(row, 5, user, format)
    sheet1.write_string(row, 6, phone, format)
    sheet1.write_string(row, 7, ip, format)
    sheet1.write_string(row, 8, entity, format)
    sheet1.write_string(row, 9, fulladdress, format)
    sheet1.write_string(row, 10, city, format)
    sheet1.write_string(row, 11, state, format)
    sheet1.write_string(row, 12, zip, format)
    sheet1.write_string(row, 13, country, format)
    sheet1.write_string(row, 14, note, format)
    sheet1.write_string(row, 15, aka, format)
    sheet1.write_string(row, 16, dob, format)
    sheet1.write_string(row, 17, gender, format)
    sheet1.write_string(row, 18, info, format)
    sheet1.write_string(row, 19, misc, format)
    sheet1.write_string(row, 20, lastname, format)
    sheet1.write_string(row, 21, firstname, format)
    sheet1.write_string(row, 22, middlename, format)
    sheet1.write_string(row, 23, friend, format)
    sheet1.write_string(row, 24, otherurls, format)
    sheet1.write_string(row, 25, otherphones, format)
    sheet1.write_string(row, 26, otheremails, format)
    sheet1.write_string(row, 27, case, format)
    sheet1.write_string(row, 28, sosfilenumber, format)
    sheet1.write_string(row, 29, president, format)
    sheet1.write_string(row, 30, sosagent, format)
    sheet1.write_string(row, 31, managers, format)
    sheet1.write_string(row, 32, dnsdomain, format)
    sheet1.write_string(row, 33, dstip, format)
    sheet1.write_string(row, 34, srcip, format)
    sheet1.write_string(row, 35, content, format)
    sheet1.write_string(row, 36, referer, format)
    sheet1.write_string(row, 37, osurl, format)
    sheet1.write_string(row, 38, titleurl, format)
    try:
        sheet1.write_string(row, 39, pagestatus, format)
    except:
        pass
    row += 1

def youtube(): # testuser = kevinrose
    print(f'{color_yellow}\n\t<<<<< youtube {color_blue}users{color_yellow} >>>>>{color_reset}')
    for user in users:    
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = ('https://www.youtube.com/%s' %(user))
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        titleurl = titleurl.replace(' - YouTube','')
        if '404' not in pagestatus:
            fullname = titleurl
            
            if fullname.lower() == user.lower():
                fullname = ''

            print(f'{color_green}{url}{color_yellow}	{fullname}{color_reset}')
            
            write_ossint(user, '4 - youtube', fullname, url, '', user, '', '', '', ''
                , city, '', '', country, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', titleurl, pagestatus)        

def usage():
    '''
    working examples of syntax
    '''
    file = sys.argv[0].split('\\')[-1]
    print(f'\nDescription: {color_green}{description}{color_reset}')
    print(f'{file} Version: {version} by {author}')
    print(f'\n    {color_yellow}insert your input into input.txt')
    print(f'\nExample:')
    print(f'    {file} -b')
    print(f'    {file} -E')
    print(f'    {file} -i')
    print(f'    {file} -l')
    print(f'    {file} -t')
    print(f'    {file} -s')
    print(f'    {file} -p')
    print(f'    {file} -U')
    print(f'    {file} -W')
    print(f'    {file} -E -i -p -U -W -I input.txt')


if __name__ == '__main__':
    main()

# <<<<<<<<<<<<<<<<<<<<<<<<<<Revision History >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
2.8.7 - fixed regex_phone, skip internet check if it's a virtual machine
2.8.6 - made the .py and .exe version dummy proof. Just double click and it runs
2.7.6 - internet checker, removed -I and -O requirement
2.8.0 - kik
"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<Future Wishlist  >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
fix dnsdomains if it ends with / , take it off
tkinter purely gui interface
.replace isn't working in several modules
instagramtwo()
create a new identity_hunt with xlsx and requests instead of urllib2

https://vimeo.com/john

https://www.fiverr.com/samanvay 

https://hubpages.com/@kevinrose

https://www.tripadvisor.com/Profile/kevinrose

"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<     notes            >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
if 1- main, then write a report of your findings.


"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<     The End        >>>>>>>>>>>>>>>>>>>>>>>>>>